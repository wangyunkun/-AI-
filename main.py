#!/usr/bin/env python3
# -*- coding: utf-8 -*-


import sys, os, json, base64, time, re, traceback, math
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple
from collections import defaultdict, Counter

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["QT_API"] = "pyqt6"

import httpx
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from PyQt6.QtCore import (Qt, QThread, pyqtSignal, QTimer, QPointF, QRectF,
                          QBuffer, QByteArray, QIODevice, QSize, QPropertyAnimation, QEasingCurve, pyqtProperty)
from PyQt6.QtGui import (QPixmap, QColor, QAction, QPainter, QPen, QFont,
                         QImage, QPainterPath, QBrush, QKeySequence, QPalette, QLinearGradient, QIcon)
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QPushButton, QLabel, QListWidget, QListWidgetItem, QSplitter,
                             QScrollArea, QFrame, QFileDialog, QProgressBar, QMessageBox, QDialog,
                             QFormLayout, QLineEdit, QComboBox, QToolBar, QSizePolicy, QTabWidget,
                             QTextEdit, QGroupBox, QGridLayout, QSpinBox, QPlainTextEdit, QDialogButtonBox,
                             QToolButton, QMenu, QInputDialog, QGraphicsView, QGraphicsScene,
                             QGraphicsPixmapItem, QGraphicsRectItem, QGraphicsEllipseItem, QGraphicsPathItem,
                             QGraphicsTextItem, QGraphicsItem, QCheckBox, QRadioButton, QButtonGroup,
                             QSlider, QGraphicsDropShadowEffect,
                             QTreeWidget, QTreeWidgetItem)

# ==================== 全局配置 ====================
CONFIG_FILE = "app_config_v4.json"
HISTORY_FILE = "inspection_history_v4.json"
STATS_FILE = "inspection_stats_v4.json"
TEMPLATE_DIR = "templates"
MAX_IMAGES = 50
EXPORT_IMG_DIR = "_export_marked"

# 主题色配置
THEME_COLORS = {
    "primary": "#1976D2",
    "secondary": "#424242",
    "success": "#4CAF50",
    "warning": "#FF9800",
    "danger": "#F44336",
    "info": "#2196F3",
    "light": "#F5F5F5",
    "dark": "#212121",
    "severe_safety": "#D32F2F",
    "general_safety": "#F57C00",
    "severe_quality": "#E64A19",
    "general_quality": "#FFA726",
    "severe_safety_bg": "#FFEBEE",
    "general_safety_bg": "#FFF3E0",
    "severe_quality_bg": "#FBE9E7",
    "general_quality_bg": "#FFF8E1",
}

# ==================== 知识库配置 ====================
KB_INDEX_FILE = "knowledge_bases.json"
KB_DIR = "knowledge_bases"

DEFAULT_BUSINESS_DATA = {
    "company_project_map": {
        "勐海县泽兴供水有限公司": ["城乡供水一体化项目", "勐海农村供水保障项目"],
        "勐海县润博水利投资有限公司": ["勐阿水库建设项目"],
        "江城县润成水利投资有限公司": ["热水河水库建设项目"],
        "澜沧县润成水利投资有限公司": ["三道箐水库建设项目"]
    },
    "company_unit_map": {
        "勐海县泽兴供水有限公司": "云南建投第二水利水电建设有限公司",
        "勐海县润博水利投资有限公司": "云南建投第二水利水电建设有限公司",
        "江城县润成水利投资有限公司": "云南建投第二水利水电建设有限公司",
        "澜沧县润成水利投资有限公司": "云南省水利水电工程有限公司"
    },
    "check_content_options": [
        "安全生产专项检查", "工程质量专项检查", "安全质量综合检查",
        "节前安全检查", "复工安全质量检查", "专项整治检查"
    ],
    "project_overview_map": {
        "勐海农村供水保障项目": "本工程位于西双版纳州勐海县，主要建设内容包括新建取水坝、输水管网及配套水厂设施，旨在解决周边5个乡镇的农村饮水安全问题，设计供水规模为2.5万吨/日。",
        "城乡供水一体化项目": "勐海县城乡供水一体化建设项目涉及勐海县城、勐遮镇、勐混镇、勐阿镇、打洛镇、勐满镇、格朗和乡、勐宋乡8个片区，覆盖现状人口28.53万人。主要建设内容为：新建3座水厂，总建设规模32000m³/d，其中县城三水厂20000m³/d，格朗和乡4000m³/d，勐混镇 8000m³/d。扩建水厂1座，勐遮镇扩容建设5000m³/d 工艺设施，扩容后总处理规模15000m³/d。利用存量水厂7座，现状总供水规模61500m³/d，其中县城一水厂10000m³/d，县城二水厂 20000m³/d，勐遮水厂10000m³/d，打洛镇曼彦水厂7500m³/d，勐阿水厂6000m³/d，勐满水厂4000m³/d，勐宋水厂4000m³/d。建设DN100-DN900输配水管网376.87km，配套建设信息化设施、阀门井、排泥阀、闸阀、入户管及其他附属设施。",
        "勐阿水库建设项目": "勐海县勐阿水库项目总投资7.645亿元。勐海县勐阿水库规模为中型，由枢纽工程、输(供)水工程和水厂工程组成，枢纽工程由大坝、溢洪道和输水(兼导流)隧洞组成。合同工期48个月。",
        "热水河水库建设项目": "江城县热水河水库工程主要由枢纽工程和输水工程两部分组成。枢纽工程主要包括拦河坝、溢洪道、导流输水隧洞等，建成后将有效缓解江城县城供水压力。江城热水河水库项目总投资5.61亿元。江城县热水河水库工程由枢纽工程和输水工程组成，合同工期为48个月。",
        "三道箐水库建设项目": "澜沧县三道箐水库位于澜沧县中北部的东河乡拉巴河上游的三道箐河上，水库工程由枢纽工程及灌区工程组成。枢纽工程主要由大坝、1～2#副坝、溢洪道、输水导流兼放空隧洞及主坝～1#副坝库岸防渗组成。水库为小（1）型水库，总库容406万m3，澜沧县三道箐水库项目总投资2.32808亿元，合同工期24个月。"
    }
}

# ==================== 规范知识库 V6.0 专业增强版 ====================
# 针对 V5.0 问题深度不足、专业度不够进行全面优化
# 新增：典型重大隐患清单、一眼识别技巧、完整规范条文

REGULATION_DATABASE_V6 = {
    "管道": {
        "role_desc": "管道与阀门工艺专家 | 30 年压力管道安装经验",
        "critical_hazards": [
            "压力管道使用非压力管道管材（如用排水管代替给水管）",
            "阀门无标识或标识错误（介质流向、压力等级）",
            "法兰垫片使用错误（石棉垫片用于高温高压）",
            "管道支吊架间距过大导致管道下垂",
            "补偿器未做预拉伸或限位措施",
        ],
        "norms": """
### GB 50242-2002《建筑给水排水及采暖工程施工质量验收规范》
**第 3.3.13 条** 法兰连接螺栓紧固后露出螺母 2-3 扣，垫片不突入管内，法兰平行度偏差不大于法兰外径的 1.5‰。
**第 3.3.15 条** 阀门安装前必须做强度和严密性试验，安装方向正确（止回阀低进高出），手轮便于操作。
**第 4.1.2 条** 给水管道必须采用与管材相适应的管件，生活给水系统管材必须符合饮用水卫生标准。
### TSG D0001-2009《压力管道安全技术监察规程》
**第 110 条** 压力管道元件必须具有特种设备制造许可证，安装前进行外观检查和几何尺寸检查。
""",
        "checklist": [
            "【一眼识别】管道颜色标识：红色 - 消防、绿色 - 给水、蓝色 - 排水、黄色 - 燃气，颜色错误立即报告",
            "【一眼识别】法兰螺栓露牙：必须露出 2-3 扣，少于 2 扣未紧固，多于 3 扣螺栓过长",
            "【一眼识别】阀门手轮方向：向上或向外为正确，向下为错误（无法操作）",
            "【一眼识别】软接头状态：自然状态为正常，拉伸或压缩超过 15% 立即报告",
            "【工艺检查】法兰平行度：用钢尺测量两侧间距，差值大于 2mm 报告",
            "【工艺检查】支吊架间距：DN50 不超 5m，DN100 不超 6m，过大导致下垂",
            "【工艺检查】焊缝外观：无裂纹、未熔合、夹渣、气孔，咬边深度不超 0.5mm",
            "【材料检查】管材标识：查看喷码，压力管道必须有 GB/T 编号和钢号",
        ],
        "must_report_if": [
            "发现管道有凹陷、裂纹、严重腐蚀",
            "发现法兰垫片外露不均匀或偏置",
            "发现阀门铭牌缺失或模糊不清",
            "发现管道支吊架锈蚀严重或脱落",
            "发现不同材质管道直接焊接无过渡件",
        ],
        "anti_hallucination": "临时封堵盲板不是缺阀门；试压用临时支撑不是支架不足；保温层保护板接缝不是裂缝。"
    },
    "电气": {
        "role_desc": "注册电气工程师 | 30 年变配电及施工现场经验",
        "critical_hazards": [
            "临时用电未采用 TN-S 接零保护系统",
            "配电箱未做重复接地或接地电阻大于 10Ω",
            "一闸多机（一个开关控制多台设备）",
            "电缆直接拖地或浸水",
            "带电体裸露无防护罩",
            "使用铜丝、铁丝代替熔断器熔体",
        ],
        "norms": """
### GB 50303-2015《建筑电气工程施工质量验收规范》
**第 12.1.1 条** 金属桥架及其支架全长应不少于 2 处与接地干线相连，非镀锌桥架连接板两端跨接铜芯接地线截面积不小于 4mm²。
**第 14.1.1 条** 箱 (盘) 内 PE 线应通过汇流排连接，严禁串联连接，PE 线截面积符合设计要求。
**第 5.1.1 条** 三相或单相交流单芯电缆不得单独穿于钢导管内，必须同穿于一管防止涡流发热。
### JGJ 46-2005《施工现场临时用电安全技术规范》
**第 8.1.3 条** 每台用电设备必须有各自专用的开关箱，严禁用同一个开关箱直接控制 2 台及以上用电设备。
**第 5.1.1 条** 临时用电工程必须采用 TN-S 接零保护系统，实行三级配电两级保护。
""",
        "checklist": [
            "【一眼识别】电线颜色：黄绿双色只能是 PE 线，用作他途立即报告",
            "【一眼识别】配电箱门：必须有跨接软铜线，缺失立即报告",
            "【一眼识别】插座接线：左零右火上接地，接反立即报告",
            "【一眼识别】电缆敷设：直接拖地、过路无保护管立即报告",
            "【工艺检查】桥架跨接：每 30-50m 一处接地点，连接板处必须有跨接线",
            "【工艺检查】箱内接线：一机一闸一漏一箱，多机共用开关立即报告",
            "【工艺检查】漏电保护：测试漏保按钮，不动作立即报告",
            "【工艺检查】电缆弯曲半径：不小于电缆外径 10 倍，过小损伤绝缘",
            "【设备检查】配电箱标识：必须有一机一闸标识、责任人、联系电话",
        ],
        "must_report_if": [
            "发现电线绝缘层破损、老化开裂",
            "发现开关箱内积尘、积水或有异物",
            "发现漏电保护器失效或拆除",
            "发现电缆接头裸露无绝缘包扎",
            "发现配电箱未上锁或箱门缺失",
        ],
        "anti_hallucination": "施工中临时接线待整理；旧规范 PE 线可能不是黄绿双色 (2000 年前)；备用回路不是故障。"
    },
    "结构": {
        "role_desc": "结构总工程师 | 30 年混凝土及钢结构经验",
        "critical_hazards": [
            "模板支撑体系立杆悬空或垫板缺失",
            "高大模板未设置扫地杆、剪刀撑",
            "钢筋主筋位置错误（梁底筋放成面筋）",
            "混凝土浇筑后出现贯穿裂缝",
            "钢结构高强螺栓未终拧或梅花头未拧掉",
            "后浇带未按方案留设或提前拆除支撑",
        ],
        "norms": """
### GB 50204-2015《混凝土结构工程施工质量验收规范》
**第 5.5.1 条** 钢筋保护层厚度：梁类构件 +10mm/-7mm，板类构件 +8mm/-5mm，合格率 90% 以上。
**第 8.3.2 条** 施工缝留设位置：柱基础顶面、梁底面、板底面，继续浇筑前凿毛清理。
### JGJ 162-2008《建筑施工模板安全技术规范》
**第 6.1.2 条** 模板支架立杆底部必须设置垫板，严禁悬空，垫板厚度不小于 50mm。
**第 6.2.4 条** 满堂模板支架四边与中间每隔四排立杆设置一道纵向剪刀撑，由底至顶连续设置。
""",
        "checklist": [
            "【一眼识别】立杆底部：悬空、无垫板、垫板破裂立即报告",
            "【一眼识别】钢筋间距：用肉眼观察，明显不均匀或露筋立即报告",
            "【一眼识别】混凝土裂缝：宽度超 0.3mm 或贯穿性裂缝立即报告",
            "【一眼识别】模板变形：鼓胀、扭曲、下沉立即报告",
            "【工艺检查】钢筋绑扎：梁柱节点核心区箍筋不得遗漏",
            "【工艺检查】保护层垫块：每平方米不少于 1 块，梅花形布置",
            "【工艺检查】模板垂直度：层高 5m 内偏差小于 6mm",
            "【工艺检查】钢结构焊缝：焊脚尺寸符合设计，无咬边、未焊透",
            "【材料检查】钢筋锈蚀：表面锈皮脱落、出现麻坑不得使用",
        ],
        "must_report_if": [
            "发现模板支撑立杆间距大于方案要求",
            "发现梁底模板下挠或支撑松动",
            "发现钢筋规格型号与设计不符",
            "发现混凝土蜂窝、孔洞、夹渣",
            "发现钢结构涂装漏涂、返锈",
        ],
        "anti_hallucination": "未抹面不是不平整；待绑扎区域钢筋散乱正常；温度裂缝 (发丝状) 不是结构裂缝。"
    },
    "机械": {
        "role_desc": "起重机械专家 | 30 年塔吊施工升降机安拆经验",
        "critical_hazards": [
            "【致命】使用非吊装机械进行吊装作业（如用挖掘机、装载机吊物）",
            "【致命】塔吊力矩限制器、起重量限制器失效或被短接",
            "【致命】施工升降机防坠安全器过期或失效",
            "【致命】吊篮安全锁失效或配重不足",
            "钢丝绳一个节距内断丝超过 10% 继续使用",
            "起重机械未经检测验收或检测不合格继续使用",
            "特种作业人员无证操作",
        ],
        "norms": """
### GB 5144-2006《塔式起重机安全规程》
**第 6.1.1 条** 塔吊必须装设力矩限制器、起重量限制器、高度限位器、幅度限位器、回转限位器，灵敏可靠。
**第 7.2.1 条** 钢丝绳报废标准：一个节距内断丝数超过总丝数 10%，有断股、死弯、压扁、绳芯挤出。
**第 10.3 条** 塔吊基础不得积水，基础周围不得挖掘，地脚螺栓紧固并有防松措施。
### GB 10055-2007《施工升降机安全规程》
**第 11.1.9 条** 防坠安全器必须在有效标定期内使用，有效期为 1 年。
### JGJ 33-2012《建筑机械使用安全技术规程》
**第 4.1.14 条** 严禁使用挖掘机、装载机、推土机等非起重机械进行吊装作业。
""",
        "checklist": [
            "【一眼识别】吊装设备：挖掘机、装载机吊物立即报告 (致命违章)",
            "【一眼识别】限位器：查看是否有线头短接、拆除，失效立即报告",
            "【一眼识别】钢丝绳：断丝、断股、死弯、压扁立即报告",
            "【一眼识别】吊钩：防脱钩装置缺失或损坏立即报告",
            "【一眼识别】配重：吊篮配重块不足或固定失效立即报告",
            "【工艺检查】标准节螺栓：用扳手检查，松动立即报告",
            "【工艺检查】附墙装置：间距符合说明书，不得焊接在脚手架上",
            "【工艺检查】基础排水：塔吊基础积水立即报告",
            "【资料检查】验收标牌：查看设备验收合格证、检测标志",
            "【人员检查】操作证：无证操作立即报告",
        ],
        "must_report_if": [
            "发现起重机械超负荷使用",
            "发现多塔作业无防碰撞措施",
            "发现塔吊回转范围内有高压线无防护",
            "发现施工升降机门联锁失效",
            "发现机械设备带病运转",
        ],
        "anti_hallucination": "停工状态吊钩无荷载正常；设备表面轻微锈迹不是缺陷；临时停放不是故障。"
    },
    "基坑": {
        "role_desc": "岩土工程师 | 30 年深基坑支护及降水经验",
        "critical_hazards": [
            "基坑开挖超过 5m 无专项方案或未按方案支护",
            "基坑边堆载超过设计荷载（堆土、堆料、机械）",
            "支护结构出现裂缝、位移、渗漏",
            "基坑降水导致周边建筑物沉降开裂",
            "上下基坑无专用通道或通道设置不合理",
            "基坑监测数据超预警值未停工",
        ],
        "norms": """
### JGJ 120-2012《建筑基坑支护技术规程》
**第 8.1.1 条** 基坑周边 1m 范围内不得堆载，3m 范围内堆载不得超过设计荷载限值。
**第 8.1.4 条** 基坑开挖过程中必须采取降排水措施，坑底不得长期浸泡。
**第 9.1.2 条** 基坑监测项目包括：支护结构位移、周边建筑物沉降、地下水位、支撑轴力。
### JGJ 59-2011《建筑施工安全检查标准》
**第 3.11.3 条** 基坑开挖深度超过 2m 必须设置 1.2m 高防护栏杆，挂密目安全网。
""",
        "checklist": [
            "【一眼识别】坑边堆载：坑边 1m 内有堆土、堆料立即报告",
            "【一眼识别】临边防护：无 1.2m 护栏或护栏损坏立即报告",
            "【一眼识别】支护裂缝：喷锚面裂缝宽度超 5mm 立即报告",
            "【一眼识别】坑底积水：大面积积水或管涌立即报告",
            "【工艺检查】放坡坡度：土质松软地区放坡不足立即报告",
            "【工艺检查】锚杆间距：符合方案要求，偏差大于 100mm 报告",
            "【工艺检查】排水沟：基坑顶截水沟、底排水沟是否畅通",
            "【工艺检查】上下通道：专用梯道宽度不小于 1m，两侧扶手",
            "【监测检查】位移观测点：是否破坏，数据是否超预警",
        ],
        "must_report_if": [
            "发现支护结构有明显位移或变形",
            "发现基坑周边地面开裂",
            "发现锚杆、土钉拔出力不足",
            "发现降水井出水量突然减少或浑浊",
            "发现基坑监测数据连续超预警值",
        ],
        "anti_hallucination": "雨后坑边少量积水及时抽排正常；支护表面轻微渗水不是渗漏；临时堆土待运不是违规。"
    },
    "消防": {
        "role_desc": "注册消防工程师 | 30 年施工现场消防管理经验",
        "critical_hazards": [
            "氧气瓶与乙炔瓶混放或间距不足 5m",
            "气瓶距离明火作业点不足 10m",
            "气瓶无防震圈、防倾倒措施",
            "动火作业无监护人、无灭火器材",
            "消防通道堵塞、消防水源不足",
            "易燃材料堆放区无消防器材",
            "工人宿舍使用大功率电器或私拉乱接",
        ],
        "norms": """
### GB 50720-2011《建设工程施工现场消防安全技术规范》
**第 5.3.7 条** 氧气瓶与乙炔瓶工作间距不小于 5m，与明火作业点距离不小于 10m，气瓶不得暴晒、不得靠近热源。
**第 6.3.1 条** 施工现场动火作业必须办理动火许可证，设专人监护，配备灭火器材。
**第 4.2.1 条** 施工现场应设置临时消防车道，宽度不小于 4m，不得占用消防车道堆放材料。
**第 5.4.3 条** 易燃易爆危险品库房与在建工程防火间距不小于 15m。
""",
        "checklist": [
            "【一眼识别】气瓶间距：氧气乙炔瓶距离小于 5m 立即报告",
            "【一眼识别】气瓶状态：横放、暴晒、无防震圈立即报告",
            "【一眼识别】动火监护：无监护人、无灭火器立即报告",
            "【一眼识别】消防通道：被材料堵塞立即报告",
            "【工艺检查】灭火器配置：每 50㎡不少于 2 具，压力指针在绿区",
            "【工艺检查】临时消防水：是否有水，水压是否足够",
            "【工艺检查】易燃物堆放：保温材料、油漆单独存放",
            "【工艺检查】宿舍消防：不得使用大功率电器、不得私拉电线",
            "【资料检查】动火证：查看是否办理、是否在有效期内",
        ],
        "must_report_if": [
            "发现乙炔瓶卧放使用",
            "发现气瓶软管破损、接头漏气",
            "发现灭火器过期或压力不足",
            "发现电焊作业无接火盆、防火毯",
            "发现消防栓无水或配件缺失",
        ],
        "anti_hallucination": "空瓶待运可横放；少量油漆当天用完可暂存；食堂用火有专人管理不是违规。"
    },
    "安全": {
        "role_desc": "注册安全工程师 | 30 年施工现场安全管理经验",
        "critical_hazards": [
            "【致命】使用非吊装机械进行吊装作业（挖掘机、装载机、汽车吊等违章吊装）",
            "高处作业 (2m 以上) 不系安全带或低挂高用",
            "安全帽未系下颌带或佩戴不合格安全帽",
            "临边洞口防护缺失或防护不牢固",
            "脚手架未满铺脚手板或探头板",
            "安全网破损、老化、未系挂",
            "交叉作业无隔离措施",
            "恶劣天气 (6 级风、暴雨) 继续高处作业",
            "起重机械作业区域无警戒、无人监护",
        ],
        "norms": """
### JGJ 59-2011《建筑施工安全检查标准》
**第 3.2.5 条** 进入施工现场必须正确佩戴安全帽，系好下颌带，安全帽必须有合格证。
**第 5.1.1 条** 高处作业 (2m 及以上) 必须系安全带，安全带必须高挂低用，挂点牢固可靠。
**第 3.13.3 条** 楼梯口、电梯井口、通道口、预留洞口必须设置防护栏杆或盖板。
### JGJ 130-2011《建筑施工扣件式钢管脚手架安全技术规范》
**第 6.2.2 条** 脚手架作业层必须满铺脚手板，不得有探头板，外侧设置 180mm 高挡脚板。
""",
        "checklist": [
            "【一眼识别】安全帽：未系下颌带立即报告，帽壳裂纹立即更换",
            "【一眼识别】安全带：2m 以上无安全带或低挂高用立即报告",
            "【一眼识别】临边防护：无 1.2m 护栏立即报告",
            "【一眼识别】洞口防护：无盖板或盖板不固定立即报告",
            "【一眼识别】脚手板：未满铺、有探头板立即报告",
            "【工艺检查】安全网：破损、老化、未系满立即报告",
            "【工艺检查】防护栏杆：上杆 1.2m、下杆 0.6m、挡脚板 180mm",
            "【工艺检查】电梯井防护：每层 (不大于 10m) 一道水平网",
            "【工艺检查】通道防护：安全通道顶部双层防护，间距 600mm",
        ],
        "must_report_if": [
            "发现安全带挂在不牢固构件上",
            "发现安全防护设施被拆除或挪作他用",
            "发现工人酒后上岗",
            "发现特种作业无证操作",
            "发现安全警示标志缺失",
        ],
        "anti_hallucination": "管理人员在安全通道内检查可短时摘帽；地面作业不强制系安全带；休息区可不戴安全帽。"
    },
    "暖通": {
        "role_desc": "暖通工程师 | 30 年通风空调及采暖经验",
        "critical_hazards": [
            "风管穿越防火分区未设防火阀",
            "排烟管道未做独立支吊架",
            "空调冷热水管道保温层破损结露",
            "风机盘管冷凝水管倒坡",
            "设备基础未做减振或减振器失效",
        ],
        "norms": """
### GB 50243-2016《通风与空调工程施工质量验收规范》
**第 4.2.1 条** 风管法兰垫片厚度 3-5mm，不得凸入管内，垫片接头不得少于 2 处。
**第 6.2.3 条** 防火阀距墙表面距离不大于 200mm，必须设独立支吊架。
**第 8.2.4 条** 冷冻水管道保温层厚度符合设计，不得有冷桥现象。
""",
        "checklist": [
            "【一眼识别】风管支吊架：间距过大 (边长≤400mm 不超 4m) 立即报告",
            "【一眼识别】保温层：破损、脱落、结露立即报告",
            "【一眼识别】防火阀：位置错误、无法操作立即报告",
            "【工艺检查】法兰垫片：不得有直缝对接，垫片不得双拼",
            "【工艺检查】管道坡度：冷凝水管坡度不小于 0.01",
            "【工艺检查】软连接：长度 150-300mm，不得扭曲",
            "【工艺检查】设备找平：水平度偏差小于 1/1000",
        ],
        "must_report_if": [
            "发现风管漏光、漏风",
            "发现阀门安装方向错误",
            "发现管道穿墙无套管",
            "发现设备运行异常噪音",
        ],
        "anti_hallucination": "测试用临时管线；调试阶段部分阀门未开启正常。"
    },
    "给排水": {
        "role_desc": "给排水工程师 | 30 年市政及建筑给排水经验",
        "critical_hazards": [
            "排水管道倒坡或坡度不足",
            "压力管道未做强度试验",
            "排水管道未做闭水试验",
            "消防管道阀门常闭未锁定",
            "给水管道与生活水源混接",
        ],
        "norms": """
### GB 50268-2008《给水排水管道工程施工及验收规范》
**第 5.3.1 条** 管道基础砂垫层厚度不小于 100mm，管道不得直接放在原状土上。
**第 9.1.1 条** 压力管道必须进行水压试验，试验压力为工作压力的 1.5 倍。
**第 9.4.1 条** 无压管道必须进行闭水试验，试验水头为上游管道内顶以上 2m。
""",
        "checklist": [
            "【一眼识别】管道坡度：排水管道倒坡立即报告",
            "【一眼识别】管道支墩：大口径管道无支墩立即报告",
            "【一眼识别】检查井：井盖破损、井内淤堵立即报告",
            "【工艺检查】管道接口：橡胶圈安装到位，无扭曲",
            "【工艺检查】地漏水封：水封深度不小于 50mm",
            "【工艺检查】管道冲洗：出水清澈无杂质",
        ],
        "must_report_if": [
            "发现管道渗漏",
            "发现阀门启闭不灵活",
            "发现管道标识缺失",
            "发现检查井内有害气体",
        ],
        "anti_hallucination": "临时排水管；施工阶段未通水正常。"
    },
    "防水": {
        "role_desc": "防水工程师 | 30 年屋面及地下防水经验",
        "critical_hazards": [
            "地下室底板防水层破损渗漏",
            "屋面卷材搭接宽度不足",
            "卫生间防水层未上翻或高度不足",
            "穿墙管防水处理不当渗漏",
            "防水保护层未及时施工导致防水层破坏",
        ],
        "norms": """
### GB 50207-2012《屋面工程质量验收规范》
**第 4.3.1 条** 卷材搭接宽度：高聚物改性沥青防水卷材短边搭接 150mm，长边搭接 100mm。
**第 5.1.3 条** 屋面女儿墙、山墙、泛水处卷材必须满粘，收头用金属压条固定。
### GB 50108-2008《地下工程防水技术规范》
**第 4.1.7 条** 防水混凝土抗渗等级不得小于 P6，施工缝必须设置止水钢板或遇水膨胀止水条。
""",
        "checklist": [
            "【一眼识别】卷材搭接：用尺量，不足 100mm 立即报告",
            "【一眼识别】屋面积水：雨后积水或排水不畅立即报告",
            "【一眼识别】渗漏痕迹：墙面水渍、发霉立即报告",
            "【工艺检查】阴阳角：必须做圆弧处理 (R=50mm)",
            "【工艺检查】附加层：管根、地漏周围 500mm 范围附加层",
            "【工艺检查】涂膜厚度：用针刺法，平均厚度符合设计",
        ],
        "must_report_if": [
            "发现防水层起鼓、开裂",
            "发现防水层裸露未保护",
            "发现变形缝漏水",
            "发现后浇带渗漏",
        ],
        "anti_hallucination": "防水层未做保护层前不能上人；施工缝处轻微潮湿不是渗漏。"
    },
    "环保": {
        "role_desc": "环境工程师 | 30 年施工现场环保管理经验",
        "critical_hazards": [
            "裸土未覆盖或覆盖不完整",
            "施工现场未设置围挡或围挡破损",
            "噪声超标 (昼间 70dB，夜间 55dB)",
            "污水直排或未经沉淀排放",
            "建筑垃圾未及时清运",
            "焚烧建筑垃圾或废弃物",
        ],
        "norms": """
### GB 12523-2011《建筑施工场界环境噪声排放标准》
**第 4.1 条** 噪声限值：昼间 70dB(A)，夜间 55dB(A)，夜间指 22:00 至次日 6:00。
### GB 50720-2011《建设工程施工现场环境与卫生标准》
**第 4.2.1 条** 施工现场主要道路必须进行硬化处理，裸露场地和集中堆放的土方应采取覆盖、固化或绿化措施。
**第 4.3.1 条** 施工现场应设置排水沟及沉淀池，污水经沉淀达标后方可排放。
""",
        "checklist": [
            "【一眼识别】裸土覆盖：绿色密目网覆盖，破损立即更换",
            "【一眼识别】围挡：高度 2.5m(市区) 或 1.8m(郊区)，破损立即修复",
            "【一眼识别】道路硬化：主要道路必须硬化，无泥泞",
            "【一眼识别】沉淀池：三级沉淀，定期清淤",
            "【工艺检查】喷淋系统：塔吊喷淋、围挡喷淋正常运行",
            "【工艺检查】噪声监测：设置噪声监测仪，数据公示",
            "【工艺检查】车辆冲洗：出入口设置洗车槽，不带泥上路",
        ],
        "must_report_if": [
            "发现扬尘污染严重",
            "发现污水直排市政管网",
            "发现夜间超时施工扰民",
            "发现垃圾焚烧",
        ],
        "anti_hallucination": "短时扬尘配合雾炮使用；雾天不是扬尘；少量生活垃圾分类存放待运。"
    },
    "水利": {
        "role_desc": "水利水电工程总工 | 30 年大坝、堤防、渠道、水闸施工经验",
        "critical_hazards": [
            "【致命】围堰填筑未按方案分层碾压或防渗措施缺失",
            "【致命】高边坡 (超过 50m) 开挖无支护或未按方案支护",
            "【致命】隧洞开挖未执行短进尺、弱爆破、强支护原则",
            "【致命】混凝土面板堆石坝面板裂缝宽度超过 0.3mm 未处理",
            "【致命】帷幕灌浆压力超标导致地层抬动变形",
            "土石围堰防渗土工膜破损或未连续铺设",
            "大坝填筑料含水率超标或铺土过厚",
            "溢洪道、消力池等泄洪建筑物混凝土蜂窝、孔洞",
            "渠道衬砌混凝土板厚度不足或防冻层缺失",
            "水闸闸门启闭机未做荷载试验或制动失灵",
            "压力钢管焊缝无损检测不合格继续使用",
            "水下作业无专项方案或潜水员无证上岗",
        ],
        "norms": """
### SL 714-2015《水利水电工程施工安全管理导则》
**第 5.2.1 条** 施工单位必须对危险性较大的单项工程编制专项施工方案，超过一定规模的必须组织专家论证。
**第 6.1.3 条** 高边坡、深基坑、隧洞开挖、围堰施工等必须设置安全监测点，定期观测并记录。
**第 7.2.4 条** 爆破作业必须执行一炮三检制度，警戒距离符合设计要求。
### SL 310-2004《水利水电工程施工质量检验与评定规程》
**第 4.2.1 条** 原材料、中间产品必须按批次进行检验，检验合格后方可使用。
**第 5.3.2 条** 混凝土试块抗压强度必须符合设计要求，合格率 100%。
**第 6.1.1 条** 单元工程质量评定分为合格和优良两个等级，不合格必须返工处理。
### SL 260-2014《水利水电工程施工测量规范》
**第 3.2.1 条** 施工控制网必须与设计单位移交的控制点进行联测，精度符合设计要求。
**第 5.1.2 条** 大坝轴线、溢洪道中心线等 main 轴线放样误差不大于 10mm。
### SL 52-2015《水利水电工程施工安全防护设施技术规范》
**第 4.2.1 条** 临边作业必须设置 1.2m 高防护栏杆，挂密目式安全网。
**第 5.3.2 条** 隧洞施工必须设置通风、照明、排水系统，有毒有害气体浓度符合标准。
**第 6.1.1 条** 水上作业必须配备救生设备，作业人员穿救生衣。
### SL 631-2012《水利水电工程单元工程施工质量验收评定标准 - 土石方工程》
**第 4.2.3 条** 土方填筑必须分层碾压，每层厚度不大于 300mm，压实度符合设计要求。
### SL 632-2012《水利水电工程单元工程施工质量验收评定标准 - 混凝土工程》
**第 5.2.1 条** 钢筋安装位置偏差：受力钢筋间距±10mm，箍筋间距±20mm。
**第 6.1.2 条** 混凝土浇筑自由下落高度不大于 2m，超过 2m 必须设置串筒或溜槽。
""",
        "checklist": [
            "【一眼识别】围堰施工：堰体分层填筑、每层厚度不大于 300mm，一次性填筑过高立即报告",
            "【一眼识别】高边坡：开挖坡比符合设计，无倒悬、无松动岩块，锚杆、锚索外露长度符合要求",
            "【一眼识别】隧洞开挖：拱架间距符合设计，锁脚锚杆不得遗漏，初喷混凝土厚度不小于 50mm",
            "【一眼识别】大坝填筑：铺土均匀、无明显粗细颗粒集中，碾压痕迹清晰",
            "【一眼识别】混凝土面板：表面平整、无贯穿裂缝，接缝止水完好无破损",
            "【一眼识别】渠道衬砌：混凝土板厚度、平整度，伸缩缝填充饱满",
            "【一眼识别】压力钢管：焊缝外观、防腐层完整性，支座锚固可靠",
            "【一眼识别】闸门安装：门槽垂直度、止水橡皮压缩量符合设计",
            "【工艺检查】土石方填筑：每层厚度、碾压遍数、压实度检测报告",
            "【工艺检查】混凝土浇筑：配合比、坍落度、试块留置、养护记录",
            "【工艺检查】帷幕灌浆：灌浆压力、浆液配比、单位吸浆量记录",
            "【工艺检查】锚杆 (索)：钻孔深度、注浆饱满度、拉拔试验报告",
            "【工艺检查】土工膜铺设：搭接宽度不小于 100mm，焊缝严密",
            "【工艺检查】止水设施：止水带位置居中，接头热熔连接牢固",
            "【资料检查】特种作业证：爆破员、潜水员、起重工持证上岗",
            "【监测检查】变形观测点：大坝沉降、边坡位移、渗流压力监测数据",
        ],
        "must_report_if": [
            "发现围堰渗漏、管涌、裂缝等险情",
            "发现高边坡有掉块、裂缝、位移迹象",
            "发现隧洞初期支护开裂、变形",
            "发现大坝填筑料含水率过大或弹簧土",
            "发现混凝土结构有贯穿裂缝、蜂窝、孔洞",
            "发现压力钢管焊缝有裂纹、未焊透",
            "发现闸门启闭异常、制动失灵",
            "发现帷幕灌浆压力异常、地表抬动",
            "发现水下作业无防护措施",
            "发现监测数据超预警值未停工处理",
        ],
        "anti_hallucination": "施工缝凿毛处理正常；临时排水沟不是永久排水；养护期混凝土表面湿润不是渗漏；隧洞施工通风管临时断开不是无通风系统。"
    }
}

ROUTER_SYSTEM_PROMPT = """
你是一名拥有 25 年经验的工程建设总监。请扫描施工现场图片，快速识别核心施工内容，指派 **3-5名** 最对口的硬核技术专家。

### 必须从以下 11 个角色中选择（严禁编造其他角色）：
1. **管道** 2. **电气** 3. **结构** 4. **机械** 5. **基坑**
6. **消防** 7. **暖通** 8. **给排水** 9. **防水** 10. **环保** 11. **水利**

### 强制规则：
1. 始终包含 **安全** 专家。
2. 如果画面模糊或无特定专业内容，仅输出 ["安全"]。
3. **看到以下任一情形，必须选派"机械"专家**：
   - 挖掘机、装载机、推土机、压路机等工程机械
   - 塔吊、施工升降机、物料提升机
   - 汽车吊、履带吊等起重机械
   - 吊篮、高处作业吊篮
   - 混凝土泵车、搅拌机
4. **看到以下任一情形，必须选派"水利"专家**：
   - 大坝、堤防、围堰施工
   - 溢洪道、消力池、泄洪洞
   - 渠道、渡槽、倒虹吸
   - 水闸、泵站、水电站
   - 隧洞开挖、高边坡支护
   - 帷幕灌浆、土工膜铺设
5. 输出必须是 JSON 字符串列表。

示例：看到挖掘机作业 → `["机械", "安全"]`
示例：看到塔吊 → `["机械", "安全", "结构"]`
示例：看到大坝填筑 → `["水利", "安全", "机械"]`
示例：看到隧洞开挖 → `["水利", "安全", "机械"]`
示例：只有工人 → `["安全"]`
"""

SPECIALIST_PROMPT_TEMPLATE = """
你现在是一名【{role}】（{role_desc}），拥有 30 年一线经验。
请对图片进行**工艺级找茬**。不要讲大道理，只找具体的**技术通病**和**违规细节**。

### 1. 核心规范依据 (必须引用)
{norms}

### 2. 你的深度检查清单 (Checklist)
请重点扫描以下细节：
{checklist}

### 3. 误判警示 (Anti-Hallucination)
{anti_hallucination}

### 4. 输出格式严格要求 (JSON)
你必须输出一个 JSON 数组，包含以下字段：
- **risk_level**: "严重安全隐患" / "一般安全隐患" / "严重质量缺陷" / "一般质量缺陷"
- **issue**: 【{role}】+ 具体描述
- **regulation**: 规范条文号
- **correction**: 整改措施
- **bbox**: [x1, y1, x2, y2]

**JSON 示例**:
[
  {{
    "risk_level": "严重质量缺陷",
    "issue": "【{role}】DN100 止回阀安装方向错误",
    "regulation": "GB 50242-2002 第 3.3.15 条",
    "correction": "拆除重装，调整阀门方向",
    "bbox": [100, 200, 300, 400],
    "confidence": 0.98
  }}
]
"""

# V6.0 专业增强版提示词模板 - 新增典型重大隐患清单和必须报告情形
SPECIALIST_PROMPT_TEMPLATE_V6 = """
你是一名【{role}】（{role_desc}），拥有 30 年一线经验。你刚检查完现场，现在要写**整改通知单**。

## 你的任务
对图片进行**工艺级找茬**，识别**具体违规事实**，不是泛泛而谈。

## 典型重大隐患清单（必须重点检查）
{critical_hazards}

⚠️ **如果发现上述任一情形，必须立即报告为"严重安全隐患"**！

## 深度检查清单（按此逐项扫描）
{checklist}

## 必须报告的情形
{must_report_if}

✅ **发现上述任一情形，直接报告，不要犹豫**！

## 核心规范依据（引用条文必须准确）
{norms}

## 误判警示（以下情况不要误报）
{anti_hallucination}

## 输出格式要求（JSON 数组）
每个问题必须包含：
- **risk_level**: "严重安全隐患" / "一般安全隐患" / "严重质量缺陷" / "一般质量缺陷"
- **issue**: 【{role}】+ 具体描述（说人话，不要套话）
- **regulation**: 规范条文号（如"GB 50242-2002 第 3.3.15 条"）
- **correction**: 整改措施（具体可执行，如"立即停工，更换合格管材"）
- **bbox**: [x1, y1, x2, y2]（问题位置坐标）
- **confidence**: 0.0-1.0（置信度，不确定给 0.6-0.7）

**输出示例**:
[
  {{
    "risk_level": "严重安全隐患",
    "issue": "【机械】使用挖掘机进行吊装作业（吊钩挂在挖掘机铲斗上吊运钢筋）",
    "regulation": "JGJ 33-2012《建筑机械使用安全技术规程》第 4.1.14 条",
    "correction": "立即停止违章作业，使用合格起重设备进行吊装",
    "bbox": [120, 150, 480, 320],
    "confidence": 0.98
  }},
  {{
    "risk_level": "一般安全隐患",
    "issue": "【电气】配电箱门未跨接软铜线",
    "regulation": "GB 50303-2015《建筑电气工程施工质量验收规范》第 12.1.1 条",
    "correction": "在配电箱门与箱体间跨接截面积不小于 4mm²的铜芯软线",
    "bbox": [200, 100, 350, 280],
    "confidence": 0.95
  }}
]

## 最后强调
1. 不要说"可能存在""疑似"，要给出明确判断
2. 不要说"建议"，要说"必须""立即"
3. 问题描述要具体，如"DN100 止回阀装反"而不是"阀门安装不规范"
4. 看到典型重大隐患，必须报告为"严重安全隐患"
"""

DEFAULT_PROMPTS_V4 = {
    "V4.0 安全质量双聚焦": "聚焦安全隐患 + 质量问题",
    "安全隐患专项": "仅识别安全隐患（忽略质量）",
    "质量问题专项": "仅识别质量问题（忽略安全）",
    "高危风险筛查": "仅识别严重安全隐患",
}

DEFAULT_PROVIDERS = {
    "阿里百炼 (Qwen-VL-Max)": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "model": "qwen-vl-max"
    },
    "阿里百炼 (Qwen2.5-VL)": {
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "model": "qwen2.5-vl-72b"
    },
    "硅基流动 (Qwen2-VL)": {
        "base_url": "https://api.siliconflow.cn/v1",
        "model": "Qwen/Qwen2-VL-72B-Instruct"
    },
    "OpenAI(GPT-4o)": {
        "base_url": "https://api.openai.com/v1",
        "model": "gpt-4o"
    },
    "自定义": {
        "base_url": "",
        "model": ""
    }
}


# ==================== 知识库管理器 ====================
class KnowledgeBaseManager:
    @staticmethod
    def _index_path():
        return KB_INDEX_FILE

    @staticmethod
    def load_index() -> List[Dict]:
        if os.path.exists(KB_INDEX_FILE):
            try:
                with open(KB_INDEX_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                pass
        return []

    @staticmethod
    def save_index(index: List[Dict]):
        try:
            with open(KB_INDEX_FILE, 'w', encoding='utf-8') as f:
                json.dump(index, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"知识库索引保存失败：{e}")

    @staticmethod
    def add(json_path: str) -> Tuple[bool, str]:
        if not os.path.exists(json_path):
            return False, f"文件不存在：{json_path}"
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            return False, f"JSON 解析失败：{e}"

        if 'toc' not in data:
            return False, "非标准 TOC-JSON 格式（缺少 toc 字段）"

        meta = data.get('meta', {})
        title = meta.get('title', os.path.basename(json_path))
        node_count = meta.get('total_nodes', 0)

        index = KnowledgeBaseManager.load_index()

        for item in index:
            if item['path'] == json_path:
                return False, f"已存在：{title}"

        index.append({
            'id': str(int(time.time() * 1000)),
            'title': title,
            'path': json_path,
            'node_count': node_count,
            'enabled': True,
            'added_at': datetime.now().isoformat()
        })
        KnowledgeBaseManager.save_index(index)
        return True, f"✅ 已导入：{title}（{node_count}个节点）"

    @staticmethod
    def remove(kb_id: str):
        index = KnowledgeBaseManager.load_index()
        index = [x for x in index if x['id'] != kb_id]
        KnowledgeBaseManager.save_index(index)

    @staticmethod
    def toggle_enabled(kb_id: str, enabled: bool):
        index = KnowledgeBaseManager.load_index()
        for item in index:
            if item['id'] == kb_id:
                item['enabled'] = enabled
        KnowledgeBaseManager.save_index(index)

    @staticmethod
    def load_all_nodes() -> List[Dict]:
        all_nodes = []
        index = KnowledgeBaseManager.load_index()
        for kb in index:
            if not kb.get('enabled', True):
                continue
            path = kb.get('path', '')
            if not os.path.exists(path):
                continue
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                nodes = []
                KnowledgeBaseManager._flatten_toc(data.get('toc', []), nodes, kb['title'])
                all_nodes.extend(nodes)
            except Exception as e:
                print(f"知识库加载失败 {path}: {e}")
        return all_nodes

    @staticmethod
    def _flatten_toc(toc_nodes: List[Dict], result: List[Dict], kb_title: str):
        for node in toc_nodes:
            result.append({
                'kb_title': kb_title,
                'id': node.get('id', ''),
                'level': node.get('level', 1),
                'title': node.get('title', ''),
                'summary': node.get('summary', ''),
                'keywords': node.get('keywords', []),
                'content_chunk': node.get('content_chunk', ''),
                'page_start': node.get('page_start', ''),
                'page_end': node.get('page_end', ''),
            })
            KnowledgeBaseManager._flatten_toc(node.get('children', []), result, kb_title)


# ==================== RAG 检索引擎 ====================
class RAGEngine:
    @staticmethod
    def search(query: str, nodes: List[Dict], top_k: int = 3) -> List[Dict]:
        if not nodes or not query:
            return []

        scored = []
        query_lower = query.lower()
        query_chars = set(query)

        for node in nodes:
            score = 0.0

            for kw in node.get('keywords', []):
                if kw and kw in query:
                    score += 3.0
                elif kw and any(c in query for c in kw if len(kw) > 1):
                    score += 0.5

            title = node.get('title', '')
            if title:
                title_chars = set(title)
                overlap = len(query_chars & title_chars) / max(len(title_chars), 1)
                score += overlap * 2.0

            summary = node.get('summary', '')
            if summary:
                summary_chars = set(summary)
                overlap = len(query_chars & summary_chars) / max(len(summary_chars), 1)
                score += overlap * 1.5

            chunk = node.get('content_chunk', '')
            if chunk:
                words = re.findall(r'[\u4e00-\u9fa5]{2,}|[a-zA-Z]{3,}', query)
                for w in words:
                    if w in chunk:
                        score += 1.0

            safety_kws = ['安全', '隐患', '危险', '防护', '违章', '临边', '高处', '用电']
            quality_kws = ['质量', '规范', '标准', '验收', '工艺', '施工', '检验']
            for kw in safety_kws + quality_kws:
                if kw in query and kw in (title + summary + chunk):
                    score += 0.8

            if score > 0.3:
                scored.append((score, node))

        scored.sort(key=lambda x: -x[0])
        results = []
        for s, node in scored[:top_k]:
            n = dict(node)
            n['_score'] = s
            results.append(n)
        return results

    @staticmethod
    def format_for_prompt(nodes: List[Dict]) -> str:
        if not nodes:
            return ""
        lines = ["\n### 📋 本公司制度要求（请结合以下条款审查并引用）："]
        for i, node in enumerate(nodes, 1):
            kb = node.get('kb_title', '企业制度')
            title = node.get('title', '')
            chunk = node.get('content_chunk', '')
            page = node.get('page_start', '')
            page_str = f"（第{page}页）" if page else ""
            lines.append(f"\n**[制度{i}]《{kb}》- {title}{page_str}**")
            if chunk:
                lines.append(f"{chunk[:300]}")
        lines.append("\n请在发现违反上述制度要求的问题时，在 regulation 字段中引用对应制度名称和条款标题。")
        return "\n".join(lines)


# ==================== 知识库管理对话框 ====================
class KnowledgeBaseDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📚 知识库管理")
        self.resize(820, 560)
        self._setup_ui()
        self._refresh_list()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(12)

        title_row = QHBoxLayout()
        lbl = QLabel("📚  企业制度知识库管理")
        lbl.setStyleSheet("font-size: 16px; font-weight: bold; color: #1976D2;")
        title_row.addWidget(lbl)
        title_row.addStretch()
        hint = QLabel("导入由「制度文件 TOC 提炼工具」生成的 JSON 文件")
        hint.setStyleSheet("font-size: 12px; color: #9E9E9E;")
        title_row.addWidget(hint)
        layout.addLayout(title_row)

        btn_row = QHBoxLayout()
        self.btn_import = QPushButton("➕  导入知识库 JSON")
        self.btn_import.setStyleSheet(
            "background:#1976D2;color:white;font-weight:bold;padding:8px 16px;"
            "border-radius:4px;border:none;")
        self.btn_import.clicked.connect(self._import_kb)

        self.btn_remove = QPushButton("🗑️  删除选中")
        self.btn_remove.setStyleSheet(
            "background:#F44336;color:white;font-weight:bold;padding:8px 16px;"
            "border-radius:4px;border:none;")
        self.btn_remove.clicked.connect(self._remove_kb)

        self.lbl_stats = QLabel("已加载 0 个知识库 · 0 个节点")
        self.lbl_stats.setStyleSheet("color: #757575; font-size: 12px;")

        btn_row.addWidget(self.btn_import)
        btn_row.addWidget(self.btn_remove)
        btn_row.addStretch()
        btn_row.addWidget(self.lbl_stats)
        layout.addLayout(btn_row)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        left = QWidget()
        left_layout = QVBoxLayout(left)
        left_layout.setContentsMargins(0, 0, 0, 0)
        lbl_list = QLabel("已注册知识库")
        lbl_list.setStyleSheet("font-weight:bold;color:#424242;margin-bottom:4px;")
        left_layout.addWidget(lbl_list)

        self.kb_list = QListWidget()
        self.kb_list.setStyleSheet("""
            QListWidget { border: 1px solid #E0E0E0; border-radius: 4px; font-size: 13px; }
            QListWidget::item { padding: 10px 8px; border-bottom: 1px solid #F5F5F5; }
            QListWidget::item:selected { background: #E3F2FD; color: #1976D2; }
        """)
        self.kb_list.currentRowChanged.connect(self._on_select)
        left_layout.addWidget(self.kb_list)

        right = QWidget()
        right_layout = QVBoxLayout(right)
        right_layout.setContentsMargins(0, 0, 0, 0)
        lbl_detail = QLabel("条款预览")
        lbl_detail.setStyleSheet("font-weight:bold;color:#424242;margin-bottom:4px;")
        right_layout.addWidget(lbl_detail)

        self.detail_tree = QTreeWidget()
        self.detail_tree.setHeaderLabels(["条款目录", "页码", "摘要"])
        self.detail_tree.setColumnWidth(0, 220)
        self.detail_tree.setColumnWidth(1, 60)
        self.detail_tree.setStyleSheet("""
            QTreeWidget { border: 1px solid #E0E0E0; border-radius: 4px; font-size: 12px; }
            QTreeWidget::item { padding: 4px; }
            QTreeWidget::item:selected { background: #E3F2FD; }
        """)
        right_layout.addWidget(self.detail_tree)

        splitter.addWidget(left)
        splitter.addWidget(right)
        splitter.setSizes([280, 520])
        layout.addWidget(splitter)

        hint2 = QLabel("💡 提示：双击列表中的知识库可切换启用/禁用状态。")
        hint2.setStyleSheet("color: #757575; font-size: 11px; padding: 4px;")
        hint2.setWordWrap(True)
        layout.addWidget(hint2)

        self.kb_list.itemDoubleClicked.connect(self._toggle_enabled)

        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(self.accept)
        layout.addWidget(btns)

    def _refresh_list(self):
        self.kb_list.clear()
        index = KnowledgeBaseManager.load_index()
        total_nodes = 0
        for kb in index:
            enabled = kb.get('enabled', True)
            icon = "✅" if enabled else "⏸"
            node_count = kb.get('node_count', 0)
            total_nodes += node_count if enabled else 0
            item = QListWidgetItem(f"{icon}  {kb['title']}  ·  {node_count}个节点")
            item.setData(Qt.ItemDataRole.UserRole, kb['id'])
            if not enabled:
                item.setForeground(QColor("#BDBDBD"))
            self.kb_list.addItem(item)

        self.lbl_stats.setText(f"已加载 {len(index)} 个知识库 · {total_nodes} 个有效节点")

    def _on_select(self, row):
        self.detail_tree.clear()
        index = KnowledgeBaseManager.load_index()
        if row < 0 or row >= len(index):
            return
        kb = index[row]
        path = kb.get('path', '')
        if not os.path.exists(path):
            self.detail_tree.addTopLevelItem(
                QTreeWidgetItem(["⚠️ 文件不存在", "", path]))
            return
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            self._build_tree(data.get('toc', []), self.detail_tree.invisibleRootItem())
            self.detail_tree.expandToDepth(1)
        except Exception as e:
            self.detail_tree.addTopLevelItem(QTreeWidgetItem([f"解析失败：{e}", "", ""]))

    def _build_tree(self, nodes, parent_item):
        level_icons = {1: "📋", 2: "📌", 3: "•"}
        for node in nodes:
            level = node.get('level', 1)
            title = node.get('title', '')
            page = str(node.get('page_start', ''))
            summary = node.get('summary', '')[:40]
            icon = level_icons.get(level, "•")
            item = QTreeWidgetItem([f"{icon} {title}", page, summary])
            if level == 1:
                item.setForeground(0, QColor("#1565C0"))
            elif level == 2:
                item.setForeground(0, QColor("#424242"))
            parent_item.addChild(item)
            self._build_tree(node.get('children', []), item)

    def _import_kb(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "选择知识库 JSON 文件", "",
            "JSON 知识库 (*.json);;所有文件 (*)"
        )
        results = []
        for path in paths:
            ok, msg = KnowledgeBaseManager.add(path)
            results.append(msg)
        if results:
            QMessageBox.information(self, "导入结果", "\n".join(results))
            self._refresh_list()

    def _remove_kb(self):
        item = self.kb_list.currentItem()
        if not item:
            QMessageBox.information(self, "提示", "请先选择要删除的知识库")
            return
        kb_id = item.data(Qt.ItemDataRole.UserRole)
        reply = QMessageBox.question(self, "确认删除",
                                     "确定删除该知识库吗？",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            KnowledgeBaseManager.remove(kb_id)
            self._refresh_list()
            self.detail_tree.clear()

    def _toggle_enabled(self, item):
        kb_id = item.data(Qt.ItemDataRole.UserRole)
        index = KnowledgeBaseManager.load_index()
        for kb in index:
            if kb['id'] == kb_id:
                new_state = not kb.get('enabled', True)
                KnowledgeBaseManager.toggle_enabled(kb_id, new_state)
                break
        self._refresh_list()


# ==================== 配置管理 ====================
class ConfigManager:
    @staticmethod
    def get_default():
        return {
            "current_provider": "阿里百炼 (Qwen-VL-Max)",
            "api_key": "",
            "last_prompt": "V4.0 安全质量双聚焦",
            "custom_provider_settings": {"base_url": "", "model": ""},
            "business_data": DEFAULT_BUSINESS_DATA,
            "prompts": DEFAULT_PROMPTS_V4,
            "provider_presets": DEFAULT_PROVIDERS,
            "max_concurrency": 2,
            "max_retries": 2,
            "request_timeout_sec": 90,
            "temperature": 0.2,
            "last_check_person": "",
            "recent_check_areas": [],
            "enable_dual_model": False,
            "secondary_provider": "OpenAI(GPT-4o)",
            "secondary_api_key": "",
            "auto_save_history": True,
            "show_confidence": True,
            "theme": "light"
        }

    @staticmethod
    def load():
        default = ConfigManager.get_default()
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    saved = json.load(f)
                for k, v in default.items():
                    if k not in saved:
                        saved[k] = v
                if "business_data" in saved:
                    for key in default["business_data"]:
                        if key not in saved["business_data"]:
                            saved["business_data"][key] = default["business_data"][key]
                return saved
            except Exception as e:
                print(f"配置加载失败：{e}")
        ConfigManager.save(default)
        return default

    @staticmethod
    def save(config):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(config, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"配置保存失败：{e}")


# ==================== 历史记录管理 ====================
class HistoryManager:
    @staticmethod
    def load():
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return {"inspections": []}
        return {"inspections": []}

    @staticmethod
    def save(data):
        try:
            with open(HISTORY_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"历史保存失败：{e}")

    @staticmethod
    def add_record(project, date, person, stats, tasks):
        history = HistoryManager.load()
        record = {
            "id": str(int(time.time() * 1000)),
            "project": project,
            "date": date,
            "person": person,
            "stats": stats,
            "task_count": len(tasks),
            "timestamp": datetime.now().isoformat()
        }
        history["inspections"].insert(0, record)
        history["inspections"] = history["inspections"][:100]
        HistoryManager.save(history)


# ==================== 统计分析管理 ====================
class StatsManager:
    @staticmethod
    def analyze_tasks(tasks):
        stats = {
            "total_images": len(tasks),
            "analyzed_images": 0,
            "total_issues": 0,
            "severe_safety": 0,
            "general_safety": 0,
            "severe_quality": 0,
            "general_quality": 0,
            "by_specialty": defaultdict(int),
            "avg_issues_per_image": 0,
            "detection_rate": 0
        }

        all_issues = []
        for task in tasks:
            if task.get("status") == "done":
                stats["analyzed_images"] += 1
                issues = task.get("edited_issues") or task.get("issues", [])
                all_issues.extend(issues)

        for issue in all_issues:
            stats["total_issues"] += 1
            level = issue.get("risk_level", "")

            if "严重安全" in level:
                stats["severe_safety"] += 1
            elif "一般安全" in level:
                stats["general_safety"] += 1
            elif "严重质量" in level:
                stats["severe_quality"] += 1
            elif "一般质量" in level:
                stats["general_quality"] += 1

            issue_text = issue.get("issue", "")
            if "【" in issue_text and "】" in issue_text:
                specialty = issue_text.split("】")[0].replace("【", "")
                stats["by_specialty"][specialty] += 1

        if stats["analyzed_images"] > 0:
            stats["avg_issues_per_image"] = round(stats["total_issues"] / stats["analyzed_images"], 1)
            stats["detection_rate"] = round((stats["analyzed_images"] / stats["total_images"]) * 100, 1)

        return stats


# ==================== JSON 解析工具 ====================
def parse_json_safe(raw: str) -> Tuple[List[Dict], Optional[str]]:
    if not raw:
        return [], "空响应"

    text = raw.strip().replace("```json", "").replace("```JSON", "").replace("```", "")
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1:
        return [], "未找到 JSON 数组"

    text = text[start:end + 1]
    text = text.replace("'", '"').replace("None", "null").replace("True", "true").replace("False", "false")

    for attempt in range(10):
        try:
            data = json.loads(text)
            break
        except json.JSONDecodeError as e:
            if "Expecting ',' delimiter" in str(e):
                text = text[:e.pos] + "," + text[e.pos:]
            elif "Expecting property name" in str(e):
                prev = text[:e.pos].rstrip()
                if prev.endswith(","):
                    comma_idx = text.rfind(",", 0, e.pos)
                    text = text[:comma_idx] + text[e.pos:]
                else:
                    break
            else:
                break
    else:
        objects = re.findall(r'\{[^{}]+\}', text)
        if objects:
            data = []
            for obj_str in objects:
                try:
                    data.append(json.loads(obj_str))
                except:
                    continue
            if data:
                return _normalize_issues(data), None
        return [], "JSON 解析失败"

    if not isinstance(data, list):
        return [], "非数组格式"

    return _normalize_issues(data), None


def _normalize_issues(data: List[Dict]) -> List[Dict]:
    result = []
    for item in data:
        if not isinstance(item, dict):
            continue

        bbox = item.get("bbox")
        if bbox and len(bbox) == 4:
            try:
                bbox = [int(float(x)) for x in bbox]
                bbox = [max(-10000, min(100000, x)) for x in bbox]
                x1, x2 = sorted([bbox[0], bbox[2]])
                y1, y2 = sorted([bbox[1], bbox[3]])
                if x2 - x1 > 1 and y2 - y1 > 1:
                    bbox = [x1, y1, x2, y2]
                else:
                    bbox = None
            except:
                bbox = None
        else:
            bbox = None

        try:
            conf = float(item.get("confidence", 0.9))
        except:
            conf = 0.9

        level = str(item.get("risk_level", ""))
        if "安全" in level:
            category = "安全隐患"
        else:
            category = "质量问题"

        result.append({
            "risk_level": level.strip(),
            "category": item.get("category", category),
            "issue": str(item.get("issue", "")).strip(),
            "regulation": str(item.get("regulation", "")).strip(),
            "correction": str(item.get("correction", "")).strip(),
            "bbox": bbox,
            "confidence": conf
        })

    return result


def calc_iou(box1, box2):
    if not box1 or not box2:
        return 0.0
    x1 = max(box1[0], box2[0])
    y1 = max(box1[1], box2[1])
    x2 = min(box1[2], box2[2])
    y2 = min(box1[3], box2[3])
    if x2 <= x1 or y2 <= y1:
        return 0.0
    inter = (x2 - x1) * (y2 - y1)
    area1 = (box1[2] - box1[0]) * (box1[3] - box1[1])
    area2 = (box2[2] - box2[0]) * (box2[3] - box2[1])
    union = area1 + area2 - inter
    return inter / union if union > 0 else 0.0


def build_specialist_prompt(role: str) -> str:
    # 使用 V6.0 专业增强版数据库
    db = REGULATION_DATABASE_V6.get(role, REGULATION_DATABASE_V6.get("安全", {}))

    # 典型重大隐患清单（强制检查项）
    critical_hazards = db.get("critical_hazards", [])
    critical_str = "\n".join([f"⚠️ {item}" for item in critical_hazards]) if critical_hazards else "无"

    # 深度检查清单
    checklist_items = db.get("checklist", [])
    checklist = "\n".join([f"🔍 {item}" for item in checklist_items])

    # 必须报告的情形
    must_report = db.get("must_report_if", [])
    must_report_str = "\n".join([f"✅ {item}" for item in must_report]) if must_report else "无"

    return SPECIALIST_PROMPT_TEMPLATE_V6.format(
        role=role,
        role_desc=db.get("role_desc", "工程专家"),
        critical_hazards=critical_str,
        checklist=checklist,
        must_report_if=must_report_str,
        norms=db.get("norms", "相关国家规范"),
        anti_hallucination=db.get("anti_hallucination", "无")
    )


# ==================== 图片导出工具 ====================
def ensure_export_dir():
    if not os.path.exists(EXPORT_IMG_DIR):
        os.makedirs(EXPORT_IMG_DIR, exist_ok=True)
    return EXPORT_IMG_DIR


def draw_on_image(img: QImage, issues: List[Dict], anns: List[Dict]) -> QImage:
    if img.isNull():
        return img

    out = img.copy()
    painter = QPainter(out)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing)

    if anns:
        for ann in anns:
            ann_type = ann.get("type")
            color = QColor(ann.get("color", "#FF0000"))
            width = int(ann.get("width", 6))

            pen = QPen(color, width)
            pen.setCapStyle(Qt.PenCapStyle.RoundCap)
            pen.setJoinStyle(Qt.PenJoinStyle.RoundJoin)
            painter.setPen(pen)
            painter.setBrush(Qt.BrushStyle.NoBrush)

            if ann_type == "rect":
                x1, y1, x2, y2 = ann.get("bbox", [0, 0, 0, 0])
                painter.drawRect(QRectF(x1, y1, x2 - x1, y2 - y1))
            elif ann_type == "ellipse":
                x1, y1, x2, y2 = ann.get("bbox", [0, 0, 0, 0])
                painter.drawEllipse(QRectF(x1, y1, x2 - x1, y2 - y1))
            elif ann_type == "arrow":
                x1, y1 = ann.get("p1", [0, 0])
                x2, y2 = ann.get("p2", [0, 0])
                painter.drawLine(QPointF(x1, y1), QPointF(x2, y2))
                angle = math.atan2(y2 - y1, x2 - x1)
                head_len = 30
                head_ang = math.radians(25)
                p1 = QPointF(x2 - head_len * math.cos(angle - head_ang), y2 - head_len * math.sin(angle - head_ang))
                p2 = QPointF(x2 - head_len * math.cos(angle + head_ang), y2 - head_len * math.sin(angle + head_ang))
                painter.drawLine(QPointF(x2, y2), p1)
                painter.drawLine(QPointF(x2, y2), p2)
            elif ann_type == "text":
                x, y = ann.get("pos", [0, 0])
                text = ann.get("text", "")
                font_size = int(ann.get("font_size", 32))
                font = QFont("SimHei", font_size, QFont.Weight.Bold)
                painter.setFont(font)
                path = QPainterPath()
                path.addText(QPointF(x, y), font, text)
                painter.setPen(QPen(QColor(255, 255, 255), 6))
                painter.drawPath(path)
                painter.setPen(QPen(color))
                painter.drawText(QPointF(x, y), text)

    painter.end()
    return out


def export_marked_image(orig_path, issues, anns, out_path):
    if not os.path.exists(orig_path):
        return False

    img = QImage(orig_path)
    if img.isNull():
        return False

    final_img = draw_on_image(img, issues, anns)
    return final_img.save(out_path, "PNG")


# ==================== Word 报告生成 ====================
class WordReportGenerator:
    @staticmethod
    def set_font(run, font_name='宋体', size=None, bold=False, color=None):
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        if size:
            run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    @staticmethod
    def replace_placeholders(doc, info):
        s_safe = str(info.get("severe_safety", 0))
        g_safe = str(info.get("general_safety", 0))
        s_qual = str(info.get("severe_quality", 0))
        g_qual = str(info.get("general_quality", 0))
        total = str(info.get("total_issues", 0))

        replacements = {
            "{{项目公司名称}}": info.get("project_company", ""),
            "{{项目名称}}": info.get("project_name", ""),
            "{{检查部位}}": info.get("check_area", ""),
            "{{检查人员}}": info.get("check_person", ""),
            "{{被检查单位}}": info.get("inspected_unit", ""),
            "{{检查内容}}": info.get("check_content", ""),
            "{{项目概况}}": info.get("project_overview", ""),
            "{{检查日期}}": info.get("check_date", ""),
            "{{整改期限}}": info.get("rectification_deadline", ""),
            "{{严重安全隐患数}}": s_safe,
            "{{一般安全隐患数}}": g_safe,
            "{{严重质量问题数}}": s_qual,
            "{{严重质量缺陷数}}": s_qual,
            "{{一般质量问题数}}": g_qual,
            "{{一般质量缺陷数}}": g_qual,
            "{{问题总数}}": total,
            "{{隐患总数}}": total
        }

        for para in doc.paragraphs:
            for key, val in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in para.text:
                                para.text = para.text.replace(key, val)

    @staticmethod
    def generate(tasks, save_path, info, template="模板.docx"):
        if os.path.exists(template):
            doc = Document(template)
        else:
            doc = Document()
            doc.add_paragraph(f"【注意】未找到模板 {template}，使用空白格式")

        WordReportGenerator.replace_placeholders(doc, info)

        valid_tasks = [t for t in tasks if t.get("status") == "done" or t.get("annotations")]

        if not valid_tasks:
            doc.add_paragraph("【提示】没有已完成分析的任务")
            doc.save(save_path)
            return

        for idx, task in enumerate(valid_tasks, 1):
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)

            p_title = cell.paragraphs[0]
            title = f"检查点位 {idx}"
            group = (task.get("meta") or {}).get("group")
            if group:
                title += f"（{group}）"
            run_title = p_title.add_run(title)
            WordReportGenerator.set_font(run_title, size=14, bold=True)

            issues = task.get("edited_issues") or task.get("issues", [])

            severe_safety, general_safety, severe_quality, general_quality = [], [], [], []
            corrections = []

            for item in issues:
                level = item.get("risk_level", "")
                category = item.get("category", "")
                issue = item.get("issue", "")
                reg = item.get("regulation", "")
                corr = item.get("correction", "")

                if not issue:
                    continue

                full_desc = issue
                if reg and reg not in ["无", "常识"]:
                    full_desc += f"（依据：{reg}）"

                if "严重安全" in level:
                    severe_safety.append(full_desc)
                elif "一般安全" in level:
                    general_safety.append(full_desc)
                elif "严重质量" in level:
                    severe_quality.append(full_desc)
                elif "一般质量" in level:
                    general_quality.append(full_desc)

                if corr:
                    corrections.append(corr)

            def add_section(label, texts, color=None):
                if not texts:
                    return
                p = cell.add_paragraph()
                run_label = p.add_run(label)
                WordReportGenerator.set_font(run_label, bold=True, size=12, color=color)
                merged = "；".join(texts) + "。"
                run_text = p.add_run(merged)
                WordReportGenerator.set_font(run_text, size=11)

            add_section("🔴 严重安全隐患：", severe_safety, RGBColor(211, 47, 47))
            add_section("🟠 一般安全隐患：", general_safety, RGBColor(245, 124, 0))
            add_section("🟡 严重质量问题：", severe_quality, RGBColor(230, 74, 25))
            add_section("🟡 一般质量问题：", general_quality, RGBColor(255, 167, 38))

            p_corr = cell.add_paragraph()
            run_label = p_corr.add_run("整改要求：")
            WordReportGenerator.set_font(run_label, bold=True, size=12)

            if corrections:
                for i, corr in enumerate(corrections, 1):
                    p = cell.add_paragraph()
                    run = p.add_run(f"{i}. {corr}")
                    WordReportGenerator.set_font(run, size=11, color=RGBColor(46, 125, 50))
            else:
                run_text = p_corr.add_run("无")
                WordReportGenerator.set_font(run_text, size=11)

            img_path = task.get("export_image_path") or task.get("path")
            if img_path and os.path.exists(img_path):
                p_img = cell.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p_img.add_run().add_picture(img_path, width=Cm(14))
                except Exception as e:
                    p_img.add_run(f"[图片加载失败：{e}]")

        doc.save(save_path)


# ==================== 分析线程 ====================
class AnalysisWorker(QThread):
    result_ready = pyqtSignal(str, dict)
    log_signal = pyqtSignal(str, str)

    PRIORITY_MAP = {
        "管道": 10, "暖通": 10, "给排水": 10,
        "电气": 10, "机械": 10,
        "结构": 10, "防水": 10, "基坑": 10,
        "消防": 9, "环保": 8,
        "安全": 5
    }

    def __init__(self, task, config, prompt, kb_nodes=None):
        super().__init__()
        self.task = task
        self.config = config
        self.prompt = prompt
        self.kb_nodes = kb_nodes or []

    def run(self):
        start_time = time.time()

        def log(msg, level="info"):
            self.log_signal.emit(msg, level)

        provider = self.config["current_provider"]
        api_key = self.config["api_key"]
        preset = self.config["provider_presets"].get(provider, {})
        base_url = preset.get("base_url")
        model = preset.get("model")

        if provider == "自定义":
            custom = self.config.get("custom_provider_settings", {})
            base_url = custom.get("base_url")
            model = custom.get("model")

        if not all([api_key, base_url, model]):
            self.result_ready.emit(self.task['id'], {
                "ok": False, "error": "配置缺失", "elapsed_sec": 0
            })
            log(f"[{self.task['name']}] 配置缺失，无法启动", "error")
            return

        img_b64 = self._compress_image(self.task["path"])
        if not img_b64:
            self.result_ready.emit(self.task['id'], {
                "ok": False, "error": "图片加载失败", "elapsed_sec": 0
            })
            log(f"[{self.task['name']}] 图片加载失败", "error")
            return

        try:
            http_client = httpx.Client(http2=False, verify=False, timeout=90)
            client = OpenAI(api_key=api_key, base_url=base_url, http_client=http_client)

            experts = ["安全"]
            log(f"🔍 [{self.task['name']}] 开始智能分诊...", "info")

            router_resp = self._call_llm(client, model, ROUTER_SYSTEM_PROMPT, img_b64, role="Router")

            if router_resp:
                try:
                    detected = json.loads(router_resp.strip())
                    if isinstance(detected, list):
                        for expert in detected:
                            if expert not in experts:
                                experts.append(expert)
                except Exception as e:
                    log(f"⚠️ Router 解析异常：{e}", "warning")

            if len(experts) == 1:
                experts.extend(["结构", "电气"])

            experts = experts[:4]
            log(f"✅ [{self.task['name']}] 专家团队：{experts}", "success")

            all_issues = []

            for role in experts:
                log(f"🔬 [{self.task['name']}] {role}专家正在检查...", "info")
                time.sleep(1.0)

                # 使用 build_specialist_prompt 函数构建 V6.0 专业增强版提示词
                specialist_prompt = build_specialist_prompt(role)

                resp = self._call_llm(client, model, specialist_prompt, img_b64, role=role)

                if resp:
                    issues, err = parse_json_safe(resp)
                    if err:
                        log(f"⚠️ {role}结果解析失败：{err}", "warning")
                    else:
                        for item in issues:
                            if not item["issue"].startswith("【"):
                                item["issue"] = f"【{role}】{item['issue']}"
                        all_issues.extend(issues)
                        log(f"    - {role} 发现 {len(issues)} 个问题", "info")

            before_cnt = len(all_issues)
            final_issues = self._deduplicate(all_issues)
            log(f"✅ [{self.task['name']}] 分析完成 (去重：{before_cnt}->{len(final_issues)})", "success")

            KB_SCORE_THRESHOLD = 2.0
            if self.kb_nodes and final_issues:
                log(f"📚 [{self.task['name']}] 正在关联企业制度条款...", "info")
                matched_count = 0
                for issue_item in final_issues:
                    issue_text = issue_item.get("issue", "")
                    hits = RAGEngine.search(issue_text, self.kb_nodes, top_k=2)
                    relevant_hits = [h for h in hits if h.get("_score", 0) >= KB_SCORE_THRESHOLD]

                    if relevant_hits:
                        issue_item["kb_source"] = relevant_hits[0].get("kb_title", "")
                        issue_item["kb_chunk"] = relevant_hits[0].get("content_chunk", "")[:200]
                        issue_item["kb_refs"] = [
                            f"《{h['kb_title']}》{h['title']}" for h in relevant_hits
                        ]
                        matched_count += 1
                    else:
                        issue_item.pop("kb_source", None)
                        issue_item.pop("kb_chunk", None)
                        issue_item.pop("kb_refs", None)

                log(f"    - 企业制度关联完成（{matched_count}/{len(final_issues)} 条问题命中）", "success")

            elapsed = round(time.time() - start_time, 2)

            self.result_ready.emit(self.task["id"], {
                "ok": True,
                "issues": final_issues,
                "elapsed_sec": elapsed,
                "provider": provider,
                "model": model
            })

        except Exception as e:
            elapsed = round(time.time() - start_time, 2)
            err_msg = str(e)
            log(f"❌ [{self.task['name']}] 线程异常：{err_msg}", "error")
            print(traceback.format_exc())
            self.result_ready.emit(self.task["id"], {
                "ok": False,
                "error": err_msg,
                "issues": [],
                "elapsed_sec": elapsed
            })

    def _compress_image(self, path):
        try:
            from PyQt6.QtGui import QImageReader
            reader = QImageReader(path)
            if not reader.canRead():
                return ""
            size = reader.size()
            if size.width() > 1536 or size.height() > 1536:
                reader.setScaledSize(size.scaled(1536, 1536, Qt.AspectRatioMode.KeepAspectRatio))
            img = reader.read()
            if img.isNull():
                return ""
            ba = QByteArray()
            buf = QBuffer(ba)
            buf.open(QIODevice.OpenModeFlag.WriteOnly)
            img.save(buf, "JPEG", 85)
            return ba.toBase64().data().decode()
        except:
            return ""

    def _call_llm(self, client, model, system_prompt, img_b64, role="Unknown"):
        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user",
                 "content": [{"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}},
                             {"type": "text", "text": "输出 JSON"}]}
            ]
            temp = 0.2 if role == "Router" else 0.3
            resp = client.chat.completions.create(model=model, messages=messages, temperature=temp)
            return resp.choices[0].message.content
        except:
            return None

    def _deduplicate(self, issues):
        if not issues:
            return []
        for item in issues:
            role = item["issue"].split("】")[0].replace("【", "")
            item["_score"] = self.PRIORITY_MAP.get(role, 5)
        issues.sort(key=lambda x: x["_score"], reverse=True)
        unique = []
        for cand in issues:
            is_dup = False
            cand_bbox = cand.get("bbox")
            if not cand_bbox:
                for exist in unique:
                    if cand["issue"][:10] == exist["issue"][:10]:
                        is_dup = True
                        break
            else:
                for exist in unique:
                    exist_bbox = exist.get("bbox")
                    if exist_bbox and calc_iou(cand_bbox, exist_bbox) > 0.4:
                        is_dup = True
                        break
            if not is_dup:
                unique.append(cand)
        return unique


# ==================== UI 组件：可编辑文字 ====================
class EditableTextItem(QGraphicsTextItem):
    def __init__(self, text, callback=None):
        super().__init__(text)
        self.callback = callback
        self.setFlags(
            QGraphicsItem.GraphicsItemFlag.ItemIsMovable |
            QGraphicsItem.GraphicsItemFlag.ItemIsSelectable |
            QGraphicsItem.GraphicsItemFlag.ItemIsFocusable
        )
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setDefaultTextColor(QColor("#FF0000"))
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.setTextInteractionFlags(Qt.TextInteractionFlag.TextEditorInteraction)
            self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsMovable, False)
            self.setFocus()
            self.setCursor(Qt.CursorShape.IBeamCursor)
            super().mouseDoubleClickEvent(event)
            if self.scene() and self.scene().views():
                self.scene().views()[0].setDragMode(QGraphicsView.DragMode.NoDrag)

    def focusOutEvent(self, event):
        self.setTextInteractionFlags(Qt.TextInteractionFlag.NoTextInteraction)
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsMovable, True)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        cursor = self.textCursor()
        cursor.clearSelection()
        self.setTextCursor(cursor)
        if self.callback:
            self.callback(self)
        if self.scene() and self.scene().views():
            view = self.scene().views()[0]
            if hasattr(view, "_tool") and view._tool == "none":
                view.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        super().focusOutEvent(event)


# ==================== UI 组件：图片标注视图 ====================
class AnnotatableImageView(QGraphicsView):
    annotation_changed = pyqtSignal()
    tool_reset = pyqtSignal()

    TOOL_NONE = "none"
    TOOL_RECT = "rect"
    TOOL_ELLIPSE = "ellipse"
    TOOL_ARROW = "arrow"
    TOOL_TEXT = "text"
    TOOL_ISSUE_TAG = "issue_tag"

    def __init__(self):
        super().__init__(QGraphicsScene())
        self._pix_item = QGraphicsPixmapItem()
        self._pix_item.setZValue(-1000)
        self._pix_item.setAcceptedMouseButtons(Qt.MouseButton.NoButton)
        self.scene().addItem(self._pix_item)
        self._tool = self.TOOL_NONE
        self._dragging = False
        self._start_pt = None
        self._temp_end_pt = None
        self._img_path = None
        self._ai_issues = []
        self._draw_color = "#FF0000"
        self._draw_width = 6
        self._img_size = (1, 1)
        self.setRenderHints(QPainter.RenderHint.Antialiasing | QPainter.RenderHint.SmoothPixmapTransform)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setMouseTracking(True)
        self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)

    def set_image(self, path):
        img = QImage(path)
        if img.isNull():
            return
        pix = QPixmap.fromImage(img)
        self._img_path = path
        self._img_size = (pix.width(), pix.height())
        self._pix_item.setPixmap(pix)
        self.scene().setSceneRect(QRectF(0, 0, pix.width(), pix.height()))
        self.resetTransform()
        self.fitInView(self.sceneRect(), Qt.AspectRatioMode.KeepAspectRatio)

    def delete_selected_items(self):
        has_deleted = False
        for item in self.scene().selectedItems():
            if item != self._pix_item:
                self.scene().removeItem(item)
                has_deleted = True
        if has_deleted:
            self.annotation_changed.emit()

    def set_tool(self, tool):
        self._tool = tool
        self._dragging = False
        if tool == self.TOOL_NONE:
            self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
            self.setCursor(Qt.CursorShape.OpenHandCursor)
        else:
            self.setDragMode(QGraphicsView.DragMode.NoDrag)
            self.setCursor(Qt.CursorShape.CrossCursor)

    def set_ai_issues(self, issues):
        self._ai_issues = issues or []

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            item = self.itemAt(event.position().toPoint())
            if isinstance(item, QGraphicsTextItem):
                self.setDragMode(QGraphicsView.DragMode.NoDrag)
                super().mousePressEvent(event)
                return
            if isinstance(item, QGraphicsItem) and item is not self._pix_item and self._tool == self.TOOL_NONE:
                self.setDragMode(QGraphicsView.DragMode.NoDrag)
                super().mousePressEvent(event)
                return
            if self._tool == self.TOOL_ISSUE_TAG:
                pos = self._to_scene_pt(event.position().toPoint())
                self._handle_issue_tag(pos)
                return
            if self._tool != self.TOOL_NONE:
                self._dragging = True
                self._start_pt = self._to_scene_pt(event.position().toPoint())
                self._temp_end_pt = self._start_pt
                return
            if self._tool == self.TOOL_NONE:
                self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._dragging and self._tool != self.TOOL_NONE:
            self._temp_end_pt = self._to_scene_pt(event.position().toPoint())
            self.viewport().update()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self._dragging:
            self._finish_drawing(event)
        super().mouseReleaseEvent(event)
        if self._tool == self.TOOL_NONE and not self.scene().focusItem():
            self.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)

    def paintEvent(self, event):
        super().paintEvent(event)
        if self._dragging and self._start_pt and self._temp_end_pt:
            painter = QPainter(self.viewport())
            painter.setPen(QPen(QColor(self._draw_color), 2, Qt.PenStyle.DashLine))
            painter.setBrush(Qt.BrushStyle.NoBrush)
            p1 = self.mapFromScene(self._start_pt)
            p2 = self.mapFromScene(self._temp_end_pt)
            x = min(p1.x(), p2.x())
            y = min(p1.y(), p2.y())
            w = abs(p1.x() - p2.x())
            h = abs(p1.y() - p2.y())
            if self._tool in (self.TOOL_RECT, self.TOOL_ELLIPSE):
                if self._tool == self.TOOL_ELLIPSE:
                    painter.drawEllipse(x, y, w, h)
                else:
                    painter.drawRect(x, y, w, h)
            elif self._tool == self.TOOL_ARROW:
                painter.drawLine(p1, p2)

    def wheelEvent(self, event):
        if event.angleDelta().y() > 0:
            self.scale(1.25, 1.25)
        else:
            self.scale(0.8, 0.8)

    def _to_scene_pt(self, view_pos):
        sp = self.mapToScene(view_pos)
        x = min(max(sp.x(), 0), self._img_size[0])
        y = min(max(sp.y(), 0), self._img_size[1])
        return QPointF(x, y)

    def _finish_drawing(self, event):
        self._dragging = False
        start = self._start_pt
        end = self._to_scene_pt(event.position().toPoint())
        self._start_pt = None
        self._temp_end_pt = None
        self.viewport().update()
        if not start or not end:
            return
        if abs(start.x() - end.x()) < 5 and abs(start.y() - end.y()) < 5:
            if self._tool == self.TOOL_TEXT:
                self._create_text(start)
            return
        data = None
        if self._tool in (self.TOOL_RECT, self.TOOL_ELLIPSE):
            x1, x2 = sorted([start.x(), end.x()])
            y1, y2 = sorted([start.y(), end.y()])
            data = {"type": self._tool, "bbox": [int(x1), int(y1), int(x2), int(y2)],
                    "color": self._draw_color, "width": self._draw_width}
        elif self._tool == self.TOOL_ARROW:
            data = {"type": "arrow", "p1": [int(start.x()), int(start.y())],
                    "p2": [int(end.x()), int(end.y())], "color": self._draw_color, "width": self._draw_width}
        elif self._tool == self.TOOL_TEXT:
            self._create_text(end)
            return
        if data:
            self._create_item_from_data(data)
            self.annotation_changed.emit()

    def _create_text(self, pos):
        dialog = QDialog(self)
        dialog.setWindowTitle("输入标注文字")
        dialog.resize(400, 150)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(15, 15, 15, 15)

        label = QLabel("请输入标注文字内容:")
        label.setStyleSheet("color: #BDBDBD; font-size: 14px; font-weight: bold; margin-bottom: 5px;")
        layout.addWidget(label)

        text_input = QLineEdit()
        text_input.setStyleSheet("""
            QLineEdit {
                color: #000000;
                background: white;
                font-size: 13px;
                padding: 8px;
                border: 2px solid #BDBDBD;
                border-radius: 4px;
            }
            QLineEdit:focus {
                border: 2px solid #2196F3;
            }
        """)
        text_input.setPlaceholderText("在此输入文字...")
        layout.addWidget(text_input)

        btn_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        btn_box.setStyleSheet("""
            QPushButton {
                background: #2196F3;
                color: white;
                font-size: 13px;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover {
                background: #1976D2;
            }
            QPushButton:pressed {
                background: #1565C0;
            }
        """)
        btn_box.accepted.connect(dialog.accept)
        btn_box.rejected.connect(dialog.reject)
        layout.addWidget(btn_box)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        text = text_input.text().strip()
        if not text:
            return

        color = "#FF0000"
        font_size = 32

        data = {"type": "text", "pos": [int(pos.x()), int(pos.y())], "text": text,
                "color": color, "font_size": font_size}
        self._create_item_from_data(data)
        self.annotation_changed.emit()

    def _handle_issue_tag(self, pos):
        if not self._ai_issues:
            QMessageBox.information(self, "提示", "当前没有 AI 识别的问题可引用")
            self.tool_reset.emit()
            return
        QTimer.singleShot(0, lambda: self._open_issue_dialog(pos))

    def _open_issue_dialog(self, pos):
        dlg = IssueSelectionDialog(self, self._ai_issues)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            data = {"type": "text", "pos": [int(pos.x()), int(pos.y())],
                    "text": dlg.selected_text, "color": dlg.selected_color, "font_size": 28}
            self._create_item_from_data(data)
            self.annotation_changed.emit()

    def _create_item_from_data(self, data):
        item_type = data.get("type")
        color = QColor(data.get("color", "#FF0000"))
        width = int(data.get("width", 6))
        item = None
        if item_type == "text":
            x, y = data.get("pos", [0, 0])
            text = data.get("text", "")
            fs = int(data.get("font_size", 28))
            item = EditableTextItem(text, callback=lambda _: self.annotation_changed.emit())
            font = QFont()
            font.setPointSize(fs)
            font.setBold(True)
            item.setFont(font)
            item.setDefaultTextColor(color)
            item.setPos(x, y)
        elif item_type == "rect":
            x1, y1, x2, y2 = data.get("bbox", [0, 0, 0, 0])
            item = QGraphicsRectItem(QRectF(x1, y1, x2 - x1, y2 - y1))
        elif item_type == "ellipse":
            x1, y1, x2, y2 = data.get("bbox", [0, 0, 0, 0])
            item = QGraphicsEllipseItem(QRectF(x1, y1, x2 - x1, y2 - y1))
        elif item_type == "arrow":
            x1, y1 = data.get("p1", [0, 0])
            x2, y2 = data.get("p2", [0, 0])
            path = QPainterPath()
            path.moveTo(x1, y1)
            path.lineTo(x2, y2)
            item = QGraphicsPathItem(path)
        if item:
            if item_type != "text":
                pen = QPen(color, width)
                pen.setCapStyle(Qt.PenCapStyle.RoundCap)
                pen.setJoinStyle(Qt.PenJoinStyle.RoundJoin)
                item.setPen(pen)
                item.setFlags(QGraphicsItem.GraphicsItemFlag.ItemIsMovable |
                              QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
            item.setData(Qt.ItemDataRole.UserRole, data.copy())
            self.scene().addItem(item)

    def set_user_annotations(self, anns):
        self.blockSignals(True)
        try:
            self.clear_annotations()
            for ann in anns:
                self._create_item_from_data(ann)
        finally:
            self.blockSignals(False)

    def get_user_annotations(self):
        self.scene().clearFocus()
        anns = []
        for item in self.scene().items():
            if item == self._pix_item:
                continue
            raw_data = item.data(Qt.ItemDataRole.UserRole)
            if not raw_data:
                continue
            data = raw_data.copy()
            if isinstance(item, EditableTextItem):
                data["text"] = item.toPlainText()
                data["pos"] = [int(item.pos().x()), int(item.pos().y())]
                data["color"] = item.defaultTextColor().name()
                data["font_size"] = item.font().pointSize() or 32
            elif isinstance(item, (QGraphicsRectItem, QGraphicsEllipseItem)):
                r = item.sceneBoundingRect()
                data["bbox"] = [int(r.left()), int(r.top()), int(r.right()), int(r.bottom())]
            elif isinstance(item, QGraphicsPathItem) and data.get("type") == "arrow":
                offset = item.pos()
                p1 = data.get("p1", [0, 0])
                p2 = data.get("p2", [0, 0])
                data["p1"] = [int(p1[0] + offset.x()), int(p1[1] + offset.y())]
                data["p2"] = [int(p2[0] + offset.x()), int(p2[1] + offset.y())]
            anns.append(data)
        return anns

    def clear_annotations(self):
        for item in list(self.scene().items()):
            if item != self._pix_item:
                self.scene().removeItem(item)
        self.annotation_changed.emit()

    def delete_selected_items(self):
        for item in self.scene().selectedItems():
            if item != self._pix_item:
                self.scene().removeItem(item)
        self.annotation_changed.emit()

    def undo(self):
        items = [i for i in self.scene().items() if i != self._pix_item]
        if items:
            self.scene().removeItem(items[0])
            self.annotation_changed.emit()


# ==================== UI 组件：问题选择对话框 ====================
class IssueSelectionDialog(QDialog):
    def __init__(self, parent, issues):
        super().__init__(parent)
        self.setWindowTitle("选择要引用的问题")
        self.resize(550, 350)
        self.selected_text = ""
        self.selected_color = "#FF0000"

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("请选择此标注关联的问题："))

        self.list_widget = QListWidget()
        self.list_widget.setStyleSheet("""
            QListWidget {
                color: #000000;
                font-size: 13px;
                background: white;
            }
            QListWidget::item {
                color: #212121;
                padding: 8px;
                border-bottom: 1px solid #F0F0F0;
            }
            QListWidget::item:selected {
                background: #E3F2FD;
                color: #000000;
            }
        """)

        for idx, item in enumerate(issues, 1):
            level = item.get("risk_level", "一般")
            category = item.get("category", "")
            desc = item.get("issue", "未知问题")
            display = f"{idx}. [{category}] {level} - {desc[:50]}"

            list_item = QListWidgetItem(display)
            list_item.setData(Qt.ItemDataRole.UserRole, desc)
            list_item.setData(Qt.ItemDataRole.UserRole + 1, level)
            list_item.setForeground(QColor("#000000"))
            self.list_widget.addItem(list_item)

        layout.addWidget(self.list_widget)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        btns.setStyleSheet("""
            QPushButton {
                background: #2196F3;
                color: white;
                font-size: 13px;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover {
                background: #1976D2;
            }
            QPushButton:pressed {
                background: #1565C0;
            }
        """)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def accept(self):
        item = self.list_widget.currentItem()
        if item:
            desc = item.data(Qt.ItemDataRole.UserRole)
            level = item.data(Qt.ItemDataRole.UserRole + 1)

            short_desc = desc[:15] + "..." if len(desc) > 15 else desc
            self.selected_text = short_desc

            if "严重安全" in level:
                self.selected_color = "#D32F2F"
            elif "一般安全" in level:
                self.selected_color = "#F57C00"
            elif "严重质量" in level:
                self.selected_color = "#E64A19"
            else:
                self.selected_color = "#FFA726"

        super().accept()


# ==================== UI 组件：问题编辑对话框 ====================
class IssueEditDialog(QDialog):
    def __init__(self, parent, item):
        super().__init__(parent)
        self.setWindowTitle("编辑问题")
        self.resize(750, 650)
        self.item = dict(item)

        main_layout = QVBoxLayout(self)

        # 基础信息表单
        form = QFormLayout()

        self.cbo_level = QComboBox()
        self.cbo_level.addItems([
            "严重安全隐患", "一般安全隐患",
            "严重质量问题", "一般质量问题"
        ])
        if self.item.get("risk_level"):
            idx = self.cbo_level.findText(self.item["risk_level"])
            if idx >= 0:
                self.cbo_level.setCurrentIndex(idx)

        self.cbo_category = QComboBox()
        self.cbo_category.addItems(["安全隐患", "质量问题"])
        if self.item.get("category"):
            idx = self.cbo_category.findText(self.item["category"])
            if idx >= 0:
                self.cbo_category.setCurrentIndex(idx)

        self.txt_issue = QPlainTextEdit()
        self.txt_issue.setPlainText(self.item.get("issue", ""))
        self.txt_issue.setMinimumHeight(80)

        self.txt_reg = QLineEdit()
        self.txt_reg.setText(self.item.get("regulation", ""))

        self.txt_corr = QPlainTextEdit()
        self.txt_corr.setPlainText(self.item.get("correction", ""))
        self.txt_corr.setMinimumHeight(80)

        form.addRow("风险等级:", self.cbo_level)
        form.addRow("问题类型:", self.cbo_category)
        form.addRow("问题描述:", self.txt_issue)
        form.addRow("规范依据:", self.txt_reg)
        form.addRow("整改建议:", self.txt_corr)

        main_layout.addLayout(form)

        # 企业制度编辑区域
        kb_group = QGroupBox("📚 企业制度引用（可人工编辑）")
        kb_layout = QVBoxLayout(kb_group)

        # 制度来源
        kb_source_layout = QHBoxLayout()
        kb_source_label = QLabel("制度来源:")
        kb_source_label.setFixedWidth(80)
        self.txt_kb_source = QLineEdit()
        self.txt_kb_source.setPlaceholderText("例如：公司安全生产管理制度")
        self.txt_kb_source.setText(self.item.get("kb_source", ""))
        kb_source_layout.addWidget(kb_source_label)
        kb_source_layout.addWidget(self.txt_kb_source)
        kb_layout.addLayout(kb_source_layout)

        # 引用条款
        kb_refs_label = QLabel("引用条款:")
        kb_refs_label.setFixedWidth(80)
        self.txt_kb_refs = QPlainTextEdit()
        self.txt_kb_refs.setPlaceholderText("每行一条引用条款")
        self.txt_kb_refs.setMinimumHeight(80)
        kb_refs_list = self.item.get("kb_refs", [])
        if isinstance(kb_refs_list, list):
            kb_refs_text = "\n".join(kb_refs_list)
        else:
            kb_refs_text = str(kb_refs_list)
        self.txt_kb_refs.setPlainText(kb_refs_text)
        kb_layout.addWidget(kb_refs_label)
        kb_layout.addWidget(self.txt_kb_refs)

        # 制度原文
        kb_chunk_label = QLabel("制度原文:")
        kb_chunk_label.setFixedWidth(80)
        self.txt_kb_chunk = QPlainTextEdit()
        self.txt_kb_chunk.setPlaceholderText("制度的具体条款内容...")
        self.txt_kb_chunk.setMinimumHeight(80)
        self.txt_kb_chunk.setPlainText(self.item.get("kb_chunk", ""))
        kb_layout.addWidget(kb_chunk_label)
        kb_layout.addWidget(self.txt_kb_chunk)

        main_layout.addWidget(kb_group)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Cancel |
            QDialogButtonBox.StandardButton.Ok
        )
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        main_layout.addWidget(btns)

    def get_value(self):
        kb_refs_text = self.txt_kb_refs.toPlainText().strip()
        if kb_refs_text:
            kb_refs_list = [line.strip() for line in kb_refs_text.split("\n") if line.strip()]
        else:
            kb_refs_list = []

        return {
            "risk_level": self.cbo_level.currentText(),
            "category": self.cbo_category.currentText(),
            "issue": self.txt_issue.toPlainText().strip(),
            "regulation": self.txt_reg.text().strip(),
            "correction": self.txt_corr.toPlainText().strip(),
            "bbox": self.item.get("bbox"),
            "confidence": self.item.get("confidence"),
            "kb_source": self.txt_kb_source.text().strip(),
            "kb_chunk": self.txt_kb_chunk.toPlainText().strip(),
            "kb_refs": kb_refs_list
        }


# ==================== UI 组件：现代化问题卡片 ====================
class ModernRiskCard(QFrame):
    edit_requested = pyqtSignal(dict)
    delete_requested = pyqtSignal(dict)

    def __init__(self, item):
        super().__init__()
        self.item = item
        self.setFrameShape(QFrame.Shape.StyledPanel)

        level = item.get("risk_level", "")
        category = item.get("category", "质量问题")

        if "严重安全" in level:
            bg_color = THEME_COLORS["severe_safety_bg"]
            border_color = THEME_COLORS["severe_safety"]
            icon = "🔴"
        elif "一般安全" in level:
            bg_color = THEME_COLORS["general_safety_bg"]
            border_color = THEME_COLORS["general_safety"]
            icon = "🟠"
        elif "严重质量" in level:
            bg_color = THEME_COLORS["severe_quality_bg"]
            border_color = THEME_COLORS["severe_quality"]
            icon = "🟡"
        else:
            bg_color = THEME_COLORS["general_quality_bg"]
            border_color = THEME_COLORS["general_quality"]
            icon = "🟡"

        self.setStyleSheet(f"""
            ModernRiskCard {{
                background-color: {bg_color};
                border-left: 5px solid {border_color};
                border-radius: 6px;
                padding: 8px;
                margin: 4px 2px;
            }}
            ModernRiskCard:hover {{
                background-color: {self._lighten(bg_color)};
            }}
        """)

        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(10)
        shadow.setColor(QColor(0, 0, 0, 30))
        shadow.setOffset(2, 2)
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)

        header = QHBoxLayout()
        lbl_level = QLabel(f"{icon} <b>{level}</b>")
        lbl_level.setStyleSheet("font-size: 13px;")
        header.addWidget(lbl_level)

        tag = QLabel(category)
        tag.setStyleSheet(f"""
            QLabel {{
                background: {border_color};
                color: white;
                padding: 3px 10px;
                border-radius: 3px;
                font-size: 11px;
                font-weight: bold;
            }}
        """)
        header.addWidget(tag)
        header.addStretch()

        # 🔴 修复：使用 lambda 捕获 item 对象
        btn_edit = QPushButton("✏️ 编辑")
        btn_edit.setFixedSize(70, 28)
        btn_edit.setStyleSheet("""
            QPushButton {
                background: white;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
            }
            QPushButton:hover {
                background: #F5F5F5;
            }
        """)
        btn_edit.clicked.connect(lambda checked, i=item: self.edit_requested.emit(i))
        header.addWidget(btn_edit)

        # 🔴 修复：使用 lambda 捕获 item 对象
        btn_del = QPushButton("🗑️ 删除")
        btn_del.setFixedSize(70, 28)
        btn_del.setStyleSheet("""
            QPushButton {
                background: white;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                color: #F44336;
            }
            QPushButton:hover {
                background: #FFEBEE;
            }
        """)
        btn_del.clicked.connect(lambda checked, i=item: self.delete_requested.emit(i))
        header.addWidget(btn_del)

        layout.addLayout(header)

        issue = item.get("issue", "")
        lbl_issue = QLabel(issue[:250] + "..." if len(issue) > 250 else issue)
        lbl_issue.setWordWrap(True)
        lbl_issue.setStyleSheet("font-size: 13px; color: #212121; margin: 8px 0;")
        layout.addWidget(lbl_issue)

        reg = item.get("regulation", "")
        if reg:
            lbl_reg = QLabel(f"📋 依据：{reg}")
            lbl_reg.setStyleSheet("font-size: 11px; color: #424242; margin: 4px 0;")
            lbl_reg.setWordWrap(True)
            layout.addWidget(lbl_reg)

        corr = item.get("correction", "")
        if corr:
            lbl_corr = QLabel(f"✅ 整改：{corr[:200]}")
            lbl_corr.setWordWrap(True)
            lbl_corr.setStyleSheet("""
                font-size: 12px; 
                color: #2E7D32; 
                font-weight: bold;
                margin: 4px 0;
                padding: 6px;
                background: rgba(76, 175, 80, 0.1);
                border-radius: 4px;
            """)
            layout.addWidget(lbl_corr)

        kb_refs = self.item.get("kb_refs", [])
        kb_source = self.item.get("kb_source", "")
        kb_chunk = self.item.get("kb_chunk", "")

        if kb_refs and kb_source:
            sep = QFrame()
            sep.setFrameShape(QFrame.Shape.HLine)
            sep.setStyleSheet("color: #BBDEFB; margin: 4px 0;")
            layout.addWidget(sep)

            kb_header = QHBoxLayout()
            kb_icon = QLabel("📚")
            kb_icon.setStyleSheet("font-size: 14px;")
            kb_title_label = QLabel("引用企业制度")
            kb_title_label.setStyleSheet("""
                font-size: 11px;
                color: #1565C0;
                font-weight: bold;
                padding: 2px 8px;
                background: #E3F2FD;
                border: 1px solid #BBDEFB;
                border-radius: 10px;
            """)
            kb_header.addWidget(kb_icon)
            kb_header.addWidget(kb_title_label)
            kb_header.addStretch()
            layout.addLayout(kb_header)

            for ref_text in kb_refs:
                lbl_ref = QLabel(f"• {ref_text}")
                lbl_ref.setWordWrap(True)
                lbl_ref.setStyleSheet("""
                    font-size: 11px;
                    color: #1565C0;
                    padding: 1px 4px 1px 20px;
                """)
                layout.addWidget(lbl_ref)

            if kb_chunk:
                lbl_chunk = QLabel(f"📋 制度原文：{kb_chunk[:120]}{'…' if len(kb_chunk) > 120 else ''}")
                lbl_chunk.setWordWrap(True)
                lbl_chunk.setStyleSheet("""
                    font-size: 11px;
                    color: #5C6BC0;
                    padding: 6px 8px;
                    background: #F3F4FF;
                    border-left: 3px solid #7986CB;
                    border-radius: 0 4px 4px 0;
                    margin: 2px 0;
                """)
                layout.addWidget(lbl_chunk)

    def _lighten(self, color):
        c = QColor(color)
        return QColor(min(255, c.red() + 10), min(255, c.green() + 10), min(255, c.blue() + 10)).name()


# ==================== UI 组件：统计卡片 ====================
class StatsCard(QFrame):
    def __init__(self, title, value, color, icon):
        super().__init__()
        self.setFixedSize(QSize(230, 80))

        self.setStyleSheet(f"""
            StatsCard {{
                background-color: {color};
                border-radius: 8px;
                color: white;
            }}
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(15, 10, 15, 10)
        layout.setSpacing(15)

        lbl_icon = QLabel(icon)
        lbl_icon.setStyleSheet("font-size: 32px; font-weight: bold; border: none; background: transparent;")
        lbl_icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(lbl_icon)

        right_container = QWidget()
        right_container.setStyleSheet("background: transparent; border: none;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 5, 0, 5)
        right_layout.setSpacing(2)

        lbl_title = QLabel(title)
        lbl_title.setStyleSheet("font-size: 13px; opacity: 0.9; font-weight: bold;")

        self.lbl_value = QLabel(str(value))
        self.lbl_value.setStyleSheet("font-size: 26px; font-weight: bold;")

        right_layout.addWidget(lbl_title)
        right_layout.addWidget(self.lbl_value)
        layout.addWidget(right_container)

    def update_value(self, value):
        self.lbl_value.setText(str(value))


# ==================== UI 组件：报告信息配置对话框 ====================
class ReportInfoDialog(QDialog):
    def __init__(self, parent, config, current_info):
        super().__init__(parent)
        self.setWindowTitle("配置报告基础信息 & 模板选择")
        self.resize(780, 500)
        self.config = config
        self.business_data = config.get("business_data", {})
        self.info = current_info.copy()

        layout = QVBoxLayout(self)

        group = QGroupBox("项目与报告配置")
        form = QGridLayout(group)
        form.setVerticalSpacing(15)
        form.setHorizontalSpacing(15)

        self.cbo_company = QComboBox()
        self.cbo_project = QComboBox()
        self.txt_unit = QLineEdit()

        self.cbo_check_type = QComboBox()
        self.cbo_check_type.setEditable(True)
        self.cbo_check_type.addItems(self.business_data.get("check_content_options", []))

        self.cbo_template = QComboBox()
        self.found_templates = self._scan_docx_files()
        self.cbo_template.addItems(self.found_templates)

        self.txt_area = QLineEdit()
        self.txt_person = QLineEdit()
        self.txt_date = QLineEdit(datetime.now().strftime("%Y-%m-%d"))
        self.txt_deadline = QLineEdit()

        self.txt_overview = QPlainTextEdit()
        self.txt_overview.setPlaceholderText("选择项目后自动填充...")
        self.txt_overview.setMaximumHeight(70)

        form.addWidget(QLabel("项目公司:"), 0, 0)
        form.addWidget(self.cbo_company, 0, 1)
        form.addWidget(QLabel("项目名称:"), 0, 2)
        form.addWidget(self.cbo_project, 0, 3)

        form.addWidget(QLabel("被检单位:"), 1, 0)
        form.addWidget(self.txt_unit, 1, 1)
        form.addWidget(QLabel("检查类型:"), 1, 2)
        form.addWidget(self.cbo_check_type, 1, 3)

        form.addWidget(QLabel("导出模板:"), 2, 0)
        form.addWidget(self.cbo_template, 2, 1)
        form.addWidget(QLabel("检查部位:"), 2, 2)
        form.addWidget(self.txt_area, 2, 3)

        form.addWidget(QLabel("检查人员:"), 3, 0)
        form.addWidget(self.txt_person, 3, 1)
        form.addWidget(QLabel("检查日期:"), 3, 2)
        form.addWidget(self.txt_date, 3, 3)

        deadline_widget = QWidget()
        h_dead = QHBoxLayout(deadline_widget)
        h_dead.setContentsMargins(0, 0, 0, 0)
        h_dead.addWidget(self.txt_deadline)
        btn_3d = QPushButton("+3 天");
        btn_3d.setFixedWidth(50)
        btn_7d = QPushButton("+7 天");
        btn_7d.setFixedWidth(50)
        btn_3d.clicked.connect(lambda: self._calc_deadline(3))
        btn_7d.clicked.connect(lambda: self._calc_deadline(7))
        h_dead.addWidget(btn_3d)
        h_dead.addWidget(btn_7d)
        form.addWidget(deadline_widget, 4, 1)

        form.addWidget(QLabel("项目概况:"), 5, 0)
        form.addWidget(self.txt_overview, 5, 1, 1, 3)

        layout.addWidget(group)

        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

        self._init_data()

        self.cbo_company.currentTextChanged.connect(self._on_company_changed)
        self.cbo_project.currentTextChanged.connect(self._on_project_changed)
        self.cbo_check_type.currentTextChanged.connect(self._auto_match_template)

    def _scan_docx_files(self):
        import glob
        files = glob.glob("*.docx")
        valid_files = []
        for f in files:
            if f.startswith("~$"):
                continue
            if f.startswith("检查报告_"):
                continue
            valid_files.append(f)

        if not valid_files:
            return ["模板.docx (未找到文件)"]
        return valid_files

    def _init_data(self):
        companies = list(self.business_data.get("company_project_map", {}).keys())
        self.cbo_company.addItems(companies)

        if self.info.get("project_company"):
            self.cbo_company.setCurrentText(self.info["project_company"])
        if self.info.get("project_name"):
            self.cbo_project.setCurrentText(self.info["project_name"])

        self.txt_unit.setText(self.info.get("inspected_unit", ""))
        self.cbo_check_type.setEditText(self.info.get("check_content", "安全质量综合检查"))
        self.txt_area.setText(self.info.get("check_area", ""))
        self.txt_person.setText(self.info.get("check_person", ""))
        self.txt_date.setText(self.info.get("check_date", datetime.now().strftime("%Y-%m-%d")))
        self.txt_deadline.setText(self.info.get("rectification_deadline", ""))
        self.txt_overview.setPlainText(self.info.get("project_overview", ""))

        last_tpl = self.info.get("template_name", "")
        if last_tpl and last_tpl in self.found_templates:
            self.cbo_template.setCurrentText(last_tpl)
        else:
            self._auto_match_template(self.cbo_check_type.currentText())

        if self.cbo_project.currentText():
            self._on_project_changed(self.cbo_project.currentText())

    def _on_company_changed(self, company_name):
        self.cbo_project.blockSignals(True)
        self.cbo_project.clear()

        projects = self.business_data.get("company_project_map", {}).get(company_name, [])
        self.cbo_project.addItems(projects)

        self.cbo_project.blockSignals(False)

        unit = self.business_data.get("company_unit_map", {}).get(company_name, "")
        self.txt_unit.setText(unit)

        if self.cbo_project.count() > 0:
            self.cbo_project.setCurrentIndex(0)
            self._on_project_changed(self.cbo_project.currentText())
        else:
            self.txt_overview.clear()

    def _on_project_changed(self, project_name):
        overview = self.business_data.get("project_overview_map", {}).get(project_name, "")
        self.txt_overview.setPlainText(overview)

    def _auto_match_template(self, check_type_text):
        if not self.found_templates:
            return

        mapping = [
            ("复工", "复工"),
            ("节前", "节前"),
            ("整治", "整治"),
            ("综合", "综合"),
            ("工程质量", "质量"),
            ("安全生产", "安全"),
            ("质量", "质量"),
            ("安全", "安全")
        ]

        target_keyword = ""

        for check_key, template_key in mapping:
            if check_key in check_type_text:
                target_keyword = template_key
                break

        if not target_keyword:
            target_keyword = "通用"

        best_match = None

        for tpl in self.found_templates:
            if target_keyword in tpl:
                best_match = tpl
                break

        if not best_match and "模板.docx" in self.found_templates:
            best_match = "模板.docx"

        if best_match:
            self.cbo_template.setCurrentText(best_match)
            print(f"检查类型：{check_type_text} -> 匹配关键词：{target_keyword} -> 选中模板：{best_match}")

    def _calc_deadline(self, days):
        try:
            base = datetime.strptime(self.txt_date.text(), "%Y-%m-%d")
            self.txt_deadline.setText((base + timedelta(days=days)).strftime("%Y-%m-%d"))
        except:
            pass

    def get_data(self):
        tpl = self.cbo_template.currentText()
        if "(未找到文件)" in tpl:
            tpl = "模板.docx"

        return {
            "project_company": self.cbo_company.currentText(),
            "project_name": self.cbo_project.currentText(),
            "inspected_unit": self.txt_unit.text(),
            "check_content": self.cbo_check_type.currentText(),
            "template_name": tpl,
            "check_area": self.txt_area.text(),
            "check_person": self.txt_person.text(),
            "check_date": self.txt_date.text(),
            "rectification_deadline": self.txt_deadline.text(),
            "project_overview": self.txt_overview.toPlainText()
        }


# ==================== UI 组件：带水波纹动画的按钮 ====================
class RippleButton(QPushButton):
    def __init__(self, text="", parent=None, color=THEME_COLORS["primary"]):
        super().__init__(text, parent)
        self.cursor_pos = QPointF()
        self.radius = 0
        self.setCursor(Qt.CursorShape.PointingHandCursor)

        self.base_color = color
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: white;
                border: 1px solid #E0E0E0;
                border-radius: 4px;
                padding: 6px 12px;
                color: #333;
                font-weight: bold;
            }}
            QPushButton:hover {{
                background-color: #F5F5F5;
                border-color: {color};
                color: {color};
            }}
        """)

        self.animation = QPropertyAnimation(self, b"radius_prop")
        self.animation.setDuration(400)
        self.animation.setEasingCurve(QEasingCurve.Type.OutQuad)

        self.radius = 0

    @pyqtProperty(float)
    def radius_prop(self):
        return self.radius

    @radius_prop.setter
    def radius_prop(self, val):
        self.radius = val
        self.update()

    def mousePressEvent(self, event):
        self.cursor_pos = event.position()
        self.radius = 0
        self.animation.stop()
        end_radius = max(self.width(), self.height()) * 1.5
        self.animation.setStartValue(0)
        self.animation.setEndValue(end_radius)
        self.animation.start()
        super().mousePressEvent(event)

    def paintEvent(self, event):
        super().paintEvent(event)
        if self.radius > 0:
            painter = QPainter(self)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)
            path = QPainterPath()
            path.addRoundedRect(QRectF(self.rect()), 4, 4)
            painter.setClipPath(path)

            brush = QBrush(QColor(self.base_color))
            color = brush.color()
            color.setAlpha(40)
            painter.setBrush(color)
            painter.setPen(Qt.PenStyle.NoPen)

            painter.drawEllipse(self.cursor_pos, self.radius, self.radius)


# ==================== 主窗口 ====================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        sys.excepthook = self._global_exception_handler

        self.config = ConfigManager.load()
        self.refresh_business_data()

        self.tasks = []
        self.current_task_id = None
        self.running_workers = {}
        self.pending_queue = []

        self.report_info = {
            "project_company": "",
            "project_name": "",
            "inspected_unit": "",
            "check_content": "安全质量综合检查",
            "template_name": "模板.docx",
            "check_area": "",
            "check_person": self.config.get("last_check_person", ""),
            "check_date": datetime.now().strftime("%Y-%m-%d"),
            "rectification_deadline": "",
            "project_overview": ""
        }

        self.init_ui()
        self.setup_shortcuts()

    def _global_exception_handler(self, exc_type, exc_value, exc_traceback):
        if issubclass(exc_type, KeyboardInterrupt):
            sys.__excepthook__(exc_type, exc_value, exc_traceback)
            return
        error_msg = ''.join(traceback.format_exception(exc_type, exc_value, exc_traceback))
        print(f"❌ 全局异常:\n{error_msg}")
        QMessageBox.critical(None, "程序错误", f"{exc_type.__name__}: {exc_value}")

    def refresh_business_data(self):
        self.business_data = self.config.get("business_data", DEFAULT_BUSINESS_DATA)

    def init_ui(self):
        self.setWindowTitle("普洱版纳区域质量安全检查助手 V4.1 专家版")
        self.resize(1450, 1000)

        self.setStyleSheet("""
            QMainWindow { background-color: #F5F5F5; }
            QToolBar { background: white; border-bottom: 1px solid #E0E0E0; spacing: 8px; padding: 8px; }
            QPushButton { padding: 6px 12px; border: 1px solid #E0E0E0; border-radius: 4px; background: white; }
            QPushButton:hover { background: #F5F5F5; }
        """)

        toolbar = QToolBar("Main")
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(20, 20))
        self.addToolBar(toolbar)

        toolbar.addWidget(QLabel("  <b>检查模式</b>: "))
        self.cbo_prompt = QComboBox()
        self.cbo_prompt.setMinimumWidth(250)
        prompts = self.config.get("prompts", DEFAULT_PROMPTS_V4)
        self.cbo_prompt.addItems(prompts.keys())
        self.cbo_prompt.setCurrentText(self.config.get("last_prompt", list(prompts.keys())[0]))
        toolbar.addWidget(self.cbo_prompt)

        toolbar.addSeparator()

        self.act_add = QAction("➕ 添加图片", self)
        self.act_run = QAction("▶ 开始分析", self)
        self.act_pause = QAction("⏸ 暂停", self)
        self.act_clear = QAction("🗑️ 清空队列", self)
        toolbar.addAction(self.act_add)
        toolbar.addAction(self.act_run)
        toolbar.addAction(self.act_pause)
        toolbar.addAction(self.act_clear)

        toolbar.addSeparator()
        self.act_info = QAction("📝 报告信息配置", self)
        self.act_info.setToolTip("配置公司、项目、人员等报告基础信息")
        toolbar.addAction(self.act_info)

        self.act_kb = QAction("📚 知识库", self)
        self.act_kb.setToolTip("导入并管理企业制度知识库")
        toolbar.addAction(self.act_kb)

        self.act_export = QAction("📄 导出报告", self)
        toolbar.addAction(self.act_export)

        empty = QWidget()
        empty.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        toolbar.addWidget(empty)
        self.act_info.triggered.connect(self.open_report_info_dialog)
        self.act_history = QAction("📚 历史", self)
        self.act_help = QAction("❓ 帮助", self)
        self.act_setting = QAction("⚙ 设置", self)
        toolbar.addAction(self.act_history)
        toolbar.addAction(self.act_help)
        toolbar.addAction(self.act_setting)

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(15, 15, 15, 15)

        stats_widget = QWidget()
        stats_widget.setStyleSheet("background: transparent;")
        stats_widget.setFixedHeight(90)
        stats_layout = QHBoxLayout(stats_widget)
        stats_layout.setSpacing(10)
        stats_layout.setContentsMargins(0, 0, 0, 5)

        self.card_severe_safety = StatsCard("严重安全隐患", 0, THEME_COLORS["severe_safety"], "🔴")
        self.card_general_safety = StatsCard("一般安全隐患", 0, THEME_COLORS["general_safety"], "🟠")
        self.card_severe_quality = StatsCard("严重质量问题", 0, "#E64A19", "🚫")
        self.card_general_quality = StatsCard("一般质量问题", 0, THEME_COLORS["general_quality"], "🟡")
        self.card_checked = StatsCard("检查图像数量", "0/0", THEME_COLORS["info"], "📸")

        stats_layout.addWidget(self.card_severe_safety)
        stats_layout.addWidget(self.card_general_safety)
        stats_layout.addWidget(self.card_severe_quality)
        stats_layout.addWidget(self.card_general_quality)
        stats_layout.addWidget(self.card_checked)
        stats_layout.addStretch()

        main_layout.addWidget(stats_widget)

        splitter_main = QSplitter(Qt.Orientation.Horizontal)

        left_widget = QWidget()
        left_widget.setStyleSheet("background: white; border-radius: 8px;")
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(10, 10, 10, 10)

        self.lbl_count = QLabel(f"<b>待检队列</b> (0/{MAX_IMAGES})")
        self.lbl_count.setStyleSheet("font-size: 14px; color: #212121;")

        self.list_widget = QListWidget()
        self.list_widget.setStyleSheet("""
             QListWidget { border: none; font-size: 13px; }
             QListWidget::item { padding: 8px; border-bottom: 1px solid #F0F0F0; }
             QListWidget::item:selected { background: #E3F2FD; color: #1976D2; }
             QListWidget::item[data-status="done"] { color: #4CAF50; }
        """)

        btn_layout = QHBoxLayout()
        self.btn_retry = QPushButton("🔄 重试失败")
        self.btn_retry.clicked.connect(self.retry_errors)
        btn_layout.addWidget(self.btn_retry)
        btn_layout.addStretch()

        left_layout.addWidget(self.lbl_count)
        left_layout.addWidget(self.list_widget)
        left_layout.addLayout(btn_layout)

        splitter_right_vertical = QSplitter(Qt.Orientation.Vertical)
        splitter_right_vertical.setHandleWidth(8)
        splitter_right_vertical.setStyleSheet("""
            QSplitter::handle:vertical {
                background: #E0E0E0;
                height: 4px;
                margin: 4px 0px;
                border-radius: 2px;
            }
            QSplitter::handle:vertical:hover { background: #BDBDBD; }
        """)

        top_container = QWidget()
        top_container.setStyleSheet("background: white; border-radius: 8px;")
        top_layout = QVBoxLayout(top_container)
        top_layout.setContentsMargins(10, 10, 10, 10)

        tool_widget = QWidget()
        tool_widget.setStyleSheet("background: #FAFAFA; border-radius: 4px; padding: 4px;")
        tool_layout = QHBoxLayout(tool_widget)
        tool_layout.setContentsMargins(5, 5, 5, 5)

        tool_layout.addWidget(QLabel("<b>标注工具:</b>"))
        self.btn_tool_rect = RippleButton("⬜ 框")
        self.btn_tool_text = RippleButton("📝 文字")
        self.btn_tool_tag = RippleButton("🏷️ 引用问题", color=THEME_COLORS["primary"])
        # 🔴 修复 1: 创建删除选中按钮
        self.btn_del_sel = RippleButton("❌ 删除选中", color=THEME_COLORS["danger"])

        self.btn_auto = RippleButton("🤖 自动标识", color=THEME_COLORS["success"])
        self.btn_save = RippleButton("💾 保存截图", color=THEME_COLORS["info"])
        self.btn_clear_anno = RippleButton("🗑️ 清空所有", color=THEME_COLORS["secondary"])

        for btn in [self.btn_tool_rect, self.btn_tool_text]:
            btn.setFixedWidth(70)
        self.btn_tool_tag.setFixedWidth(90)
        self.btn_auto.setStyleSheet("background: #4CAF50; color: white; font-weight: bold;")

        tool_layout.addWidget(self.btn_tool_rect)
        tool_layout.addWidget(self.btn_tool_text)
        tool_layout.addWidget(self.btn_tool_tag)
        tool_layout.addWidget(self.btn_del_sel)

        self.btn_change_color = RippleButton("🎨 调整文字颜色", color=THEME_COLORS["info"])
        self.btn_change_color.setFixedWidth(180)
        self.btn_change_color.setToolTip("调整选中文字的颜色")
        tool_layout.addWidget(self.btn_change_color)

        self.btn_resize_text = RippleButton("📏 调整文字大小", color=THEME_COLORS["info"])
        self.btn_resize_text.setFixedWidth(180)
        self.btn_resize_text.setToolTip("调整选中文字的字号")
        tool_layout.addWidget(self.btn_resize_text)

        tool_layout.addStretch()
        tool_layout.addWidget(self.btn_auto)
        tool_layout.addWidget(self.btn_save)
        tool_layout.addWidget(self.btn_clear_anno)

        top_layout.addWidget(tool_widget)

        self.image_view = AnnotatableImageView()
        self.image_view.setStyleSheet("border: 1px solid #E0E0E0; background: #333; border-radius: 4px;")
        top_layout.addWidget(self.image_view)

        bottom_container = QWidget()
        bottom_container.setStyleSheet("background: white; border-radius: 8px;")
        bottom_layout = QVBoxLayout(bottom_container)
        bottom_layout.setContentsMargins(10, 5, 10, 5)

        result_label = QLabel("<b>识别结果</b>")
        result_label.setStyleSheet("font-size: 13px; color: #616161; margin-bottom: 5px;")
        bottom_layout.addWidget(result_label)

        self.result_container = QWidget()
        self.result_layout = QVBoxLayout(self.result_container)
        self.result_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.result_layout.setSpacing(6)
        self.result_layout.setContentsMargins(4, 4, 4, 4)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(self.result_container)
        scroll.setStyleSheet("""
            QScrollArea { border: 1px solid #EEEEEE; border-radius: 4px; background: #FAFAFA; }
            QScrollBar:vertical { width: 10px; }
        """)
        bottom_layout.addWidget(scroll)

        splitter_right_vertical.addWidget(top_container)
        splitter_right_vertical.addWidget(bottom_container)
        splitter_right_vertical.setSizes([800, 250])
        splitter_right_vertical.setStretchFactor(0, 1)

        splitter_main.addWidget(left_widget)
        splitter_main.addWidget(splitter_right_vertical)
        splitter_main.setSizes([280, 1200])

        main_layout.addWidget(splitter_main)

        log_group = QGroupBox("运行日志")
        log_group.setFixedHeight(150)
        log_layout = QVBoxLayout(log_group)
        log_layout.setContentsMargins(5, 5, 5, 5)

        self.txt_log = QTextEdit()
        self.txt_log.setReadOnly(True)
        self.txt_log.setStyleSheet("""
                    QTextEdit {
                        background-color: #2b2b2b;
                        color: #e0e0e0;
                        font-family: Consolas, Monaco, monospace;
                        font-size: 12px;
                        border: none;
                        border-radius: 4px;
                    }
                """)
        log_layout.addWidget(self.txt_log)

        main_layout.addWidget(log_group)

        self.status_bar = self.statusBar()
        self.status_bar.setStyleSheet("background: white; color: #757575;")
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setFixedWidth(200)
        self.status_bar.addPermanentWidget(self.progress_bar)

        # 6. 信号连接
        self.act_add.triggered.connect(self.add_files)
        self.act_run.triggered.connect(self.start_analysis)
        self.act_pause.triggered.connect(self.pause_analysis)
        self.act_clear.triggered.connect(self.clear_queue)
        self.act_export.triggered.connect(self.export_word)
        self.act_history.triggered.connect(self.show_history)
        self.act_setting.triggered.connect(self.open_settings)
        self.act_help.triggered.connect(self.show_help)
        self.act_kb.triggered.connect(self.open_knowledge_base_dialog)

        self.cbo_prompt.currentTextChanged.connect(self.save_prompt_selection)
        self.list_widget.itemClicked.connect(self.on_item_clicked)

        self.btn_tool_rect.clicked.connect(lambda: self.image_view.set_tool("rect"))
        self.btn_tool_text.clicked.connect(lambda: self.image_view.set_tool("text"))
        self.btn_tool_tag.clicked.connect(lambda: self.image_view.set_tool("issue_tag"))
        # 🔴 修复 1: 连接删除选中按钮信号
        self.btn_del_sel.clicked.connect(lambda: self.image_view.delete_selected_items())

        self.btn_auto.clicked.connect(self.auto_annotate_current)
        self.btn_save.clicked.connect(self.save_marked_image)
        self.btn_clear_anno.clicked.connect(self.image_view.clear_annotations)
        self.btn_change_color.clicked.connect(self.change_selected_text_color)
        self.btn_resize_text.clicked.connect(self.resize_selected_text)

        self.image_view.annotation_changed.connect(self.on_annotation_changed)
        self.image_view.tool_reset.connect(lambda: self.image_view.set_tool("none"))

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Delete:
            self.image_view.delete_selected_items()
        else:
            super().keyPressEvent(event)

    def open_report_info_dialog(self):
        dlg = ReportInfoDialog(self, self.config, self.report_info)

        if dlg.exec() == QDialog.DialogCode.Accepted:
            self.report_info = dlg.get_data()

            if self.report_info.get("check_person"):
                self.config["last_check_person"] = self.report_info["check_person"]
                ConfigManager.save(self.config)

            p_name = self.report_info.get('project_name', '未命名项目')
            self.status_bar.showMessage(f"✅ 报告信息已更新：{p_name}", 3000)

    def setup_shortcuts(self):
        QAction("添加", self, shortcut=QKeySequence("Ctrl+O"), triggered=self.add_files)
        QAction("分析", self, shortcut=QKeySequence("F5"), triggered=self.start_analysis)
        QAction("导出", self, shortcut=QKeySequence("Ctrl+E"), triggered=lambda: self.export_word("模板.docx"))

    def save_prompt_selection(self, text):
        if text:
            self.config["last_prompt"] = text
            ConfigManager.save(self.config)

    def update_stats(self):
        stats = StatsManager.analyze_tasks(self.tasks)

        self.card_severe_safety.update_value(stats["severe_safety"])
        self.card_general_safety.update_value(stats["general_safety"])
        self.card_severe_quality.update_value(stats["severe_quality"])
        self.card_general_quality.update_value(stats["general_quality"])
        self.card_checked.update_value(f"{stats['analyzed_images']}/{stats['total_images']}")

    def log(self, message, level="info"):
        timestamp = datetime.now().strftime("[%H:%M:%S]")

        if level == "error":
            color = "#FF5252"
            icon = "❌"
        elif level == "warning":
            color = "#FFD740"
            icon = "⚠️"
        elif level == "success":
            color = "#69F0AE"
            icon = "✅"
        else:
            color = "#E0E0E0"
            icon = "ℹ️"

        html = f'<span style="color:#808080">{timestamp}</span> <span style="color:{color}">{icon} {message}</span>'
        self.txt_log.append(html)

        scrollbar = self.txt_log.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

        self.status_bar.showMessage(f"{icon} {message}", 3000)

    def add_files(self):
        current_count = len(self.tasks)
        if current_count >= MAX_IMAGES:
            QMessageBox.warning(self, "数量限制", f"单次最多 {MAX_IMAGES} 张")
            return

        remaining = MAX_IMAGES - current_count
        paths, _ = QFileDialog.getOpenFileNames(
            self, f"选择图片 (还能选{remaining}张)", "",
            "Images (*.jpg *.png *.jpeg)"
        )

        if not paths:
            return

        if len(paths) > remaining:
            QMessageBox.warning(self, "超限", f"截取前{remaining}张")
            paths = paths[:remaining]

        for path in paths:
            task_id = str(time.time()) + os.path.basename(path)
            task = {
                "id": task_id, "path": path, "name": os.path.basename(path),
                "status": "waiting", "issues": [], "edited_issues": None,
                "error": None, "elapsed_sec": None, "meta": {},
                "annotations": [], "export_image_path": None
            }
            self.tasks.append(task)

            item = QListWidgetItem(f"📷 {os.path.basename(path)}")
            item.setData(Qt.ItemDataRole.UserRole, task_id)
            self.list_widget.addItem(item)

        self.lbl_count.setText(f"<b>待检队列</b> ({len(self.tasks)}/{MAX_IMAGES})")
        self.update_stats()

    def start_analysis(self):
        if not self.config.get("api_key"):
            QMessageBox.warning(self, "配置缺失", "请先在设置中配置 API Key")
            return

        waiting = [t for t in self.tasks if t['status'] in ['waiting', 'error']]
        if not waiting:
            self.status_bar.showMessage("没有待处理任务")
            return

        for t in waiting:
            if t["id"] not in self.pending_queue and t["id"] not in self.running_workers:
                self.pending_queue.append(t["id"])
                t["status"] = "queued"

        self.progress_bar.setVisible(True)
        self._kick_scheduler()

    def pause_analysis(self):
        self.pending_queue.clear()
        for t in self.tasks:
            if t["status"] == "queued":
                t["status"] = "waiting"
        self.status_bar.showMessage("已暂停")

    def _kick_scheduler(self):
        max_conc = int(self.config.get("max_concurrency", 2))

        while len(self.running_workers) < max_conc and self.pending_queue:
            task_id = self.pending_queue.pop(0)
            task = next((t for t in self.tasks if t['id'] == task_id), None)
            if not task:
                continue

            task["status"] = "analyzing"
            self.update_list_status(task_id, "⏳")

            kb_nodes = KnowledgeBaseManager.load_all_nodes()
            if kb_nodes:
                self.log(f"📚 已加载企业制度知识库 {len(kb_nodes)} 个节点，将进行 RAG 增强", "success")

            worker = AnalysisWorker(task, self.config, "", kb_nodes=kb_nodes)
            worker.result_ready.connect(self.on_worker_done)
            worker.log_signal.connect(self.log)

            self.running_workers[task_id] = worker
            worker.start()

        total = len([t for t in self.tasks if t["status"] in ["queued", "analyzing", "done"]])
        done = len([t for t in self.tasks if t["status"] == "done"])
        if total > 0:
            self.progress_bar.setValue(int(done / total * 100))

        if not self.running_workers and not self.pending_queue:
            self.status_bar.showMessage("✅ 分析完成")
            self.progress_bar.setValue(100)
            QTimer.singleShot(2000, lambda: self.progress_bar.setVisible(False))
            if self.config.get("auto_save_history"):
                self.save_to_history()

    def on_worker_done(self, task_id, result):
        task = next((t for t in self.tasks if t['id'] == task_id), None)
        if task:
            task["elapsed_sec"] = result.get("elapsed_sec")

            if result.get("ok"):
                task['status'] = 'done'
                task['issues'] = result.get("issues", [])
                task["error"] = None

                severe_count = sum(1 for i in task['issues'] if "严重" in i.get("risk_level", ""))
                if severe_count > 0:
                    self.update_list_status(task_id, "🔴")
                elif len(task['issues']) > 0:
                    self.update_list_status(task_id, "🟡")
                else:
                    self.update_list_status(task_id, "✅")
            else:
                task['status'] = 'error'
                task['issues'] = []
                task["error"] = result.get("error")
                self.update_list_status(task_id, "❌")

            if self.current_task_id == task_id:
                QTimer.singleShot(50, lambda: self.render_result(task))

        if task_id in self.running_workers:
            worker = self.running_workers.pop(task_id)
            try:
                worker.result_ready.disconnect()
            except:
                pass
            worker.quit()
            worker.wait(1000)
            worker.deleteLater()

        self.update_stats()
        self._kick_scheduler()
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.data(Qt.ItemDataRole.UserRole) == task_id:
                task = next((t for t in self.tasks if t['id'] == task_id), None)
                if task:
                    item.setText(f"✅ {task['name']}")
                    item.setForeground(QColor("#4CAF50"))

    def on_item_clicked(self, item):
        self.current_task_id = item.data(Qt.ItemDataRole.UserRole)
        task = next((t for t in self.tasks if t['id'] == self.current_task_id), None)
        if task:
            self.render_result(task)

    def render_result(self, task):
        widgets_to_delete = []
        while self.result_layout.count():
            item = self.result_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widgets_to_delete.append(widget)

        for widget in widgets_to_delete:
            try:
                widget.blockSignals(True)
                if isinstance(widget, ModernRiskCard):
                    widget.edit_requested.disconnect()
                    widget.delete_requested.disconnect()
            except:
                pass
            widget.hide()
            widget.setParent(None)
            widget.deleteLater()

        QApplication.processEvents()

        img_path = task.get("path", "")
        if img_path and os.path.exists(img_path):
            if self.image_view._img_path != img_path:
                self.image_view.set_image(img_path)

        issues = task.get("edited_issues") if task.get("edited_issues") is not None else task.get("issues", [])
        self.image_view.set_ai_issues(issues)
        self.image_view.set_user_annotations(task.get("annotations", []) or [])

        if task['status'] == 'done':
            if issues:
                for item in issues:
                    card = ModernRiskCard(item)
                    card.edit_requested.connect(self.edit_issue)
                    card.delete_requested.connect(self.delete_issue)
                    self.result_layout.addWidget(card)
            else:
                lbl_empty = QLabel("✅ 未发现明显安全隐患或质量问题")
                lbl_empty.setStyleSheet("""
                    font-size: 14px; 
                    color: #4CAF50; 
                    padding: 20px;
                    background: #E8F5E9;
                    border-radius: 8px;
                    text-align: center;
                """)
                lbl_empty.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self.result_layout.addWidget(lbl_empty)
        elif task['status'] == 'analyzing':
            lbl_loading = QLabel("⏳ 正在分析中...")
            lbl_loading.setStyleSheet("font-size: 14px; color: #757575; padding: 20px;")
            self.result_layout.addWidget(lbl_loading)
        elif task['status'] == 'error':
            lbl_error = QLabel(f"❌ 分析失败：{task.get('error', '未知错误')}")
            lbl_error.setStyleSheet("font-size: 14px; color: #F44336; padding: 20px;")
            lbl_error.setWordWrap(True)
            self.result_layout.addWidget(lbl_error)

    # 🔴 修复 2: 修改 edit_issue 方法，使用内容匹配代替 ID 匹配
    def edit_issue(self, item):
        task = self._current_task()
        if not task or task.get("status") != "done":
            return

        issues = task.get("edited_issues") if task.get("edited_issues") is not None else task.get("issues", [])

        dlg = IssueEditDialog(self, item)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_item = dlg.get_value()

            if task.get("edited_issues") is None:
                task["edited_issues"] = [dict(x) for x in issues]

            # 🔴 修复：使用问题描述匹配而不是 ID 匹配
            for i, x in enumerate(task["edited_issues"]):
                if (x.get("issue") == item.get("issue") or
                        (x.get("risk_level") == item.get("risk_level") and
                         x.get("issue") == item.get("issue"))):
                    task["edited_issues"][i] = new_item
                    break

            task["export_image_path"] = None
            QTimer.singleShot(100, lambda: self.render_result(task))

    # 🔴 修复 3: 修改 delete_issue 方法，使用内容匹配代替 ID 匹配
    def delete_issue(self, item):
        sender_card = self.sender()
        if sender_card and isinstance(sender_card, ModernRiskCard):
            try:
                sender_card.blockSignals(True)
                sender_card.edit_requested.disconnect()
                sender_card.delete_requested.disconnect()
            except:
                pass

        task = self._current_task()
        if task:
            issues = task.get("edited_issues") if task.get("edited_issues") is not None else task.get("issues", [])
            if task.get("edited_issues") is None:
                task["edited_issues"] = [dict(x) for x in issues]

            # 🔴 修复：使用问题描述匹配而不是 ID 匹配
            task["edited_issues"] = [
                x for x in task["edited_issues"]
                if not (x.get("issue") == item.get("issue") and
                        x.get("risk_level") == item.get("risk_level"))
            ]
            task["export_image_path"] = None

            self.image_view.set_ai_issues(task["edited_issues"])
            QTimer.singleShot(150, lambda: self.render_result(task))
            self.update_stats()
            self.status_bar.showMessage("已删除该问题", 2000)

    def on_annotation_changed(self):
        task = self._current_task()
        if task:
            task["annotations"] = self.image_view.get_user_annotations()

    def auto_annotate_current(self):
        task = self._current_task()
        if not task or task.get("status") != "done":
            QMessageBox.warning(self, "提示", "请先完成 AI 分析")
            return

        issues = task.get("edited_issues") if task.get("edited_issues") is not None else task.get("issues", [])
        if not issues:
            QMessageBox.information(self, "提示", "未检测到问题")
            return

        count = 0
        for idx, item in enumerate(issues, 1):
            bbox = item.get("bbox")
            if bbox and len(bbox) == 4:
                cx = (bbox[0] + bbox[2]) / 2
                cy = (bbox[1] + bbox[3]) / 2

                text = item.get("issue", "")
                level = item.get("risk_level", "")

                import re
                text = re.sub(r'【[^】]+】', '', text)

                remove_words = [
                    "存在", "发现", "有", "未", "没有", "缺少", "应", "需要",
                    "的问题", "的情况", "的现象", "问题", "情况", "现象"
                ]
                for word in remove_words:
                    text = text.replace(word, "")

                text = re.sub(r'[，。、；：！？\s]+', '', text)

                if len(text) > 10:
                    keywords = ["不符", "不足", "不当", "未接", "未设", "缺失", "松动", "破损"]
                    for kw in keywords:
                        if kw in text:
                            pos = text.find(kw) + len(kw)
                            if 6 <= pos <= 12:
                                text = text[:pos]
                                break

                    if len(text) > 10:
                        text = text[:10]

                text = f"{idx}.{text}"

                if "严重安全" in level:
                    color = THEME_COLORS["severe_safety"]
                elif "一般安全" in level:
                    color = THEME_COLORS["general_safety"]
                elif "严重质量" in level:
                    color = THEME_COLORS["severe_quality"]
                else:
                    color = THEME_COLORS["general_quality"]

                font_size = 32

                new_anno = {"type": "text", "pos": [int(cx), int(cy)], "text": text,
                            "color": color, "width": 4, "font_size": font_size}
                self.image_view._create_item_from_data(new_anno)
                count += 1

        if count > 0:
            task["annotations"] = self.image_view.get_user_annotations()
            self.status_bar.showMessage(f"✅ 成功自动标识{count}处", 3000)

    def save_marked_image(self):
        task = self._current_task()
        if not task:
            return
        if not os.path.exists(task.get("path", "")):
            QMessageBox.warning(self, "失败", "图片不存在")
            return

        issues = task.get("edited_issues") if task.get("edited_issues") is not None else task.get("issues", [])
        anns = task.get("annotations", []) or []

        ensure_export_dir()
        base_name = os.path.splitext(os.path.basename(task["path"]))[0]
        out_path = os.path.join(EXPORT_IMG_DIR, f"{base_name}_{task['id'][-6:]}.png")

        ok = export_marked_image(task["path"], issues, anns, out_path)

        if ok:
            task["export_image_path"] = out_path
            self.status_bar.showMessage(f"✅ 已保存：{out_path}", 3000)
        else:
            QMessageBox.warning(self, "失败", "生成失败")

    def change_selected_text_color(self):
        selected_items = self.image_view.scene().selectedItems()

        text_items = [item for item in selected_items if isinstance(item, (EditableTextItem, QGraphicsTextItem))]

        if not text_items:
            QMessageBox.information(self, "提示",
                                    "请先选中要调整颜色的文字标注\n\n💡 提示：点击文字可选中，按住 Ctrl 可多选")
            return

        color_dialog = QDialog(self)
        color_dialog.setWindowTitle("选择颜色")
        color_dialog.resize(350, 200)

        layout = QVBoxLayout(color_dialog)
        layout.setContentsMargins(15, 15, 15, 15)

        info_label = QLabel(f"<b>已选中 {len(text_items)} 个文字标注</b>")
        info_label.setStyleSheet("color: #000000; font-size: 14px; margin-bottom: 10px;")
        layout.addWidget(info_label)

        color_label = QLabel("选择新颜色:")
        color_label.setStyleSheet("color: #000000; font-size: 13px; font-weight: bold; margin-bottom: 5px;")
        layout.addWidget(color_label)

        color_combo = QComboBox()
        color_combo.setStyleSheet("""
            QComboBox {
                font-size: 13px;
                padding: 5px;
                border: 2px solid #BDBDBD;
                border-radius: 4px;
            }
        """)
        colors = [
            ("🔴 红色", "#FF0000"),
            ("🟠 橙色", "#FF9800"),
            ("🟡 黄色", "#FFC107"),
            ("🟢 绿色", "#4CAF50"),
            ("🔵 蓝色", "#2196F3"),
            ("🟣 紫色", "#9C27B0"),
            ("⚫ 黑色", "#000000"),
            ("⚪ 白色", "#FFFFFF"),
        ]

        for name, code in colors:
            color_combo.addItem(name, code)

        layout.addWidget(color_combo)

        btn_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        btn_box.setStyleSheet("""
            QPushButton {
                background: #2196F3;
                color: white;
                font-size: 13px;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
                margin-top: 10px;
            }
            QPushButton:hover {
                background: #1976D2;
            }
            QPushButton:pressed {
                background: #1565C0;
            }
        """)
        btn_box.accepted.connect(color_dialog.accept)
        btn_box.rejected.connect(color_dialog.reject)
        layout.addWidget(btn_box)

        if color_dialog.exec() != QDialog.DialogCode.Accepted:
            return

        new_color_code = color_combo.currentData()
        color_name = color_combo.currentText()
        new_color = QColor(new_color_code)

        for item in text_items:
            item.setDefaultTextColor(new_color)

            data = item.data(Qt.ItemDataRole.UserRole)
            if data:
                data["color"] = new_color_code
                item.setData(Qt.ItemDataRole.UserRole, data)

            if hasattr(item, 'update'):
                item.update()

        task = self._current_task()
        if task:
            task["annotations"] = self.image_view.get_user_annotations()

        self.status_bar.showMessage(f"✅ 已调整 {len(text_items)} 个文字的颜色为{color_name}", 3000)
        self.image_view.annotation_changed.emit()

    def resize_selected_text(self):
        selected_items = self.image_view.scene().selectedItems()

        text_items = [item for item in selected_items if isinstance(item, (EditableTextItem, QGraphicsTextItem))]

        if not text_items:
            QMessageBox.information(self, "提示",
                                    "请先选中要调整字号的文字标注\n\n💡 提示：点击文字可选中，按住 Ctrl 可多选")
            return

        font_dialog = QDialog(self)
        font_dialog.setWindowTitle("调整字号")
        font_dialog.resize(350, 200)

        layout = QVBoxLayout(font_dialog)
        layout.setContentsMargins(15, 15, 15, 15)

        info_label = QLabel(f"<b>已选中 {len(text_items)} 个文字标注</b>")
        info_label.setStyleSheet("color: #000000; font-size: 14px; margin-bottom: 10px;")
        layout.addWidget(info_label)

        size_label = QLabel("选择新字号:")
        size_label.setStyleSheet("color: #000000; font-size: 13px; font-weight: bold; margin-bottom: 5px;")
        layout.addWidget(size_label)

        size_spin = QSpinBox()
        size_spin.setRange(10, 60)
        size_spin.setValue(32)
        size_spin.setSuffix("px")
        size_spin.setStyleSheet("""
            QSpinBox {
                font-size: 13px;
                padding: 5px;
                border: 2px solid #BDBDBD;
                border-radius: 4px;
            }
        """)
        layout.addWidget(size_spin)

        btn_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        btn_box.setStyleSheet("""
            QPushButton {
                background: #2196F3;
                color: white;
                font-size: 13px;
                font-weight: bold;
                padding: 6px 20px;
                border: none;
                border-radius: 4px;
                min-width: 80px;
                margin-top: 10px;
            }
            QPushButton:hover {
                background: #1976D2;
            }
            QPushButton:pressed {
                background: #1565C0;
            }
        """)
        btn_box.accepted.connect(font_dialog.accept)
        btn_box.rejected.connect(font_dialog.reject)
        layout.addWidget(btn_box)

        if font_dialog.exec() != QDialog.DialogCode.Accepted:
            return

        new_font_size = size_spin.value()

        for item in text_items:
            font = item.font()
            font.setPointSize(new_font_size)
            item.setFont(font)

            data = item.data(Qt.ItemDataRole.UserRole)
            if data:
                data["font_size"] = new_font_size
                item.setData(Qt.ItemDataRole.UserRole, data)

            if hasattr(item, 'update'):
                item.update()

        task = self._current_task()
        if task:
            task["annotations"] = self.image_view.get_user_annotations()

        self.status_bar.showMessage(f"✅ 已调整 {len(text_items)} 个文字的字号为 {new_font_size}px", 3000)
        self.image_view.annotation_changed.emit()

    def export_word(self):
        ensure_export_dir()

        valid_tasks = [t for t in self.tasks if t.get("status") == "done" or t.get("annotations")]
        if not valid_tasks:
            QMessageBox.warning(self, "提示", "没有可导出的任务。")
            return

        abs_export_dir = os.path.abspath(EXPORT_IMG_DIR)
        for t in self.tasks:
            if t not in valid_tasks:
                continue
            if not os.path.exists(t.get("path", "")):
                continue

            issues = t.get("edited_issues") if t.get("edited_issues") is not None else t.get("issues", [])
            anns = t.get("annotations", []) or []

            base_name = os.path.splitext(os.path.basename(t["path"]))[0]
            safe_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '_', '-'))
            out_path = os.path.join(abs_export_dir, f"{safe_name}_{t['id'][-6:]}.png")

            if export_marked_image(t["path"], issues, anns, out_path):
                t["export_image_path"] = out_path

        stats = StatsManager.analyze_tasks(self.tasks)
        final_info = self.report_info.copy()

        target_template = final_info.get("template_name", "模板.docx")
        if not os.path.exists(target_template):
            pass

        final_info.update({
            "severe_safety": str(stats["severe_safety"]),
            "general_safety": str(stats["general_safety"]),
            "severe_quality": str(stats["severe_quality"]),
            "general_quality": str(stats["general_quality"]),
            "total_issues": str(stats["total_issues"])
        })

        if not final_info["project_name"]:
            final_info["project_name"] = "项目名称"

        default_name = f"检查报告_{final_info['project_name']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path, _ = QFileDialog.getSaveFileName(self, "保存报告", default_name, "Word Files (*.docx)")

        if path:
            try:
                WordReportGenerator.generate(self.tasks, path, final_info, target_template)
                self.status_bar.showMessage(f"✅ 报告已生成：{path}", 5000)
            except Exception as e:
                import traceback
                QMessageBox.critical(self, "导出失败", f"生成失败:\n{e}\n{traceback.format_exc()}")

    def clear_queue(self):
        if self.running_workers:
            QMessageBox.warning(self, "警告", "任务正在分析中")
            return
        reply = QMessageBox.question(self, '确认', '确定清空所有任务吗？',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.tasks.clear()
            self.list_widget.clear()
            self.lbl_count.setText(f"<b>待检队列</b> (0/{MAX_IMAGES})")
            self.update_stats()
            while self.result_layout.count():
                self.result_layout.takeAt(0).widget().deleteLater()

    def retry_errors(self):
        error_tasks = [t for t in self.tasks if t["status"] == "error"]
        if not error_tasks:
            self.status_bar.showMessage("没有失败任务")
            return
        for t in error_tasks:
            t["status"] = "waiting"
            t["error"] = None
        self.status_bar.showMessage(f"已重置{len(error_tasks)}个任务")
        self.start_analysis()

    def save_to_history(self):
        done_tasks = [t for t in self.tasks if t.get("status") == "done"]
        if not done_tasks:
            return
        stats = StatsManager.analyze_tasks(self.tasks)
        HistoryManager.add_record("项目", datetime.now().strftime("%Y-%m-%d"),
                                  self.config.get("last_check_person", ""), stats, done_tasks)

    def show_history(self):
        history = HistoryManager.load()
        records = history.get("inspections", [])
        if not records:
            QMessageBox.information(self, "历史记录", "暂无历史记录")
            return
        dlg = QDialog(self)
        dlg.setWindowTitle("检查历史")
        dlg.resize(800, 500)
        layout = QVBoxLayout(dlg)
        list_widget = QListWidget()
        for record in records:
            stats = record.get("stats", {})
            text = f"{record['date']} | {record['project']} | {record['person']} | " \
                   f"严重安全:{stats.get('severe_safety', 0)} 质量问题:{stats.get('severe_quality', 0) + stats.get('general_quality', 0)}"
            list_widget.addItem(text)
        layout.addWidget(list_widget)
        btn_close = QPushButton("关闭")
        btn_close.clicked.connect(dlg.accept)
        layout.addWidget(btn_close)
        dlg.exec()

    def open_settings(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("系统设置")
        dlg.resize(700, 400)
        layout = QVBoxLayout(dlg)
        form = QFormLayout()

        presets = self.config.get("provider_presets", DEFAULT_PROVIDERS)
        cbo_prov = QComboBox()
        cbo_prov.addItems(presets.keys())
        cbo_prov.setCurrentText(self.config.get("current_provider", list(presets.keys())[0]))

        txt_key = QLineEdit(self.config.get("api_key", ""))
        txt_key.setEchoMode(QLineEdit.EchoMode.Password)

        form.addRow("模型厂商:", cbo_prov)
        form.addRow("API Key:", txt_key)
        layout.addLayout(form)

        btn_save = QPushButton("保存配置")
        btn_save.setStyleSheet("background: #2196F3; color: white; font-weight: bold; padding: 10px;")

        def save_all():
            self.config["current_provider"] = cbo_prov.currentText()
            self.config["api_key"] = txt_key.text().strip()
            ConfigManager.save(self.config)
            dlg.accept()
            self.status_bar.showMessage("✅ 配置已保存", 3000)

        btn_save.clicked.connect(save_all)
        layout.addWidget(btn_save)
        dlg.exec()

    def show_help(self):
        help_text = """
<h3>V4.0 终极版 - 聚焦安全质量</h3>

<h4>核心特性</h4>
<ul>
<li><b>聚焦重点</b>: 专注安全隐患和质量问题，取消文明施工</li>
<li><b>智能分诊</b>: Router 自动识别并指派 2-4 名专家</li>
<li><b>实时统计</b>: 顶部彩色卡片实时显示各类问题数量</li>
<li><b>分类显示</b>: 🔴严重安全 🟠一般安全 🟡质量问题</li>
</ul>

<h4>快捷键</h4>
<ul>
<li><b>Ctrl+O</b>: 添加图片</li>
<li><b>F5</b>: 开始分析</li>
<li><b>Ctrl+E</b>: 导出报告</li>
</ul>

<h4>操作流程</h4>
<ol>
<li>⚙ 设置 → 配置 API Key</li>
<li>➕ 添加图片 → 选择施工现场照片</li>
<li>▶ 开始分析 → 观察顶部统计卡片更新</li>
<li>查看问题卡片 → 编辑/删除</li>
<li>🤖 自动标识 → 生成序号标注</li>
<li>📄 导出报告</li>
</ol>

<p><b>注意</b>: AI 识别结果仅供参考，请人工复核。</p>
        """
        QMessageBox.information(self, "帮助", help_text)

    def update_list_status(self, task_id, icon):
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.data(Qt.ItemDataRole.UserRole) == task_id:
                task = next((t for t in self.tasks if t['id'] == task_id), None)
                if task:
                    item.setText(f"{icon} {task['name']}")

                    if icon == "✅":
                        item.setForeground(QColor("#4CAF50"))
                    elif icon == "❌":
                        item.setForeground(QColor("#F44336"))
                    elif icon == "⏳":
                        item.setForeground(QColor("#2196F3"))
                    else:
                        item.setForeground(QColor("#212121"))

    def _current_task(self):
        if not self.current_task_id:
            return None
        return next((t for t in self.tasks if t['id'] == self.current_task_id), None)

    def open_knowledge_base_dialog(self):
        dlg = KnowledgeBaseDialog(self)
        dlg.exec()
        index = KnowledgeBaseManager.load_index()
        enabled = [x for x in index if x.get('enabled', True)]
        if enabled:
            names = "、".join([x['title'][:10] for x in enabled[:3]])
            suffix = f"等{len(enabled)}个" if len(enabled) > 3 else f"{len(enabled)}个"
            self.status_bar.showMessage(
                f"📚 知识库已就绪：{names}...共{suffix}，下次分析将自动 RAG 引用", 5000)
        else:
            self.status_bar.showMessage("📚 知识库为空，分析时仅使用内置规范", 3000)


# ==================== 主函数 ====================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
